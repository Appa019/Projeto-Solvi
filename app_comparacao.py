"""
📚 Document Comparator - Aplicação Streamlit
Compara dois arquivos (PDF ou Word) e gera relatório de diferenças
Versão que ignora deslocamentos e foca apenas em alterações reais de conteúdo
"""

import streamlit as st
import fitz  # PyMuPDF
import difflib
import pandas as pd
import io
from datetime import datetime
import base64
from typing import List, Tuple, Dict, Optional, Set
import logging
from pathlib import Path
import tempfile
import os
import re

# Importações condicionais para Word
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configuração da página
st.set_page_config(
    page_title="Comparador de PDFs - Solvi",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS customizado para destacar filtros e parágrafos
st.markdown("""
<style>
    /* Estilo para o header com logo */
    .header-container {
        display: flex;
        align-items: center;
        justify-content: space-between;
        background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%);
        padding: 20px 30px;
        border-radius: 15px;
        margin-bottom: 30px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
    }
    
    .header-text {
        color: white;
        flex: 1;
    }
    
    .header-title {
        font-size: 2.5em;
        font-weight: bold;
        margin: 0;
        background: linear-gradient(45deg, #ffffff, #e0e7ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .header-subtitle {
        font-size: 1.1em;
        margin: 5px 0 0 0;
        opacity: 0.9;
    }
    
    .logo-container {
        background: white;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    .logo-container img {
        max-height: 60px;
        max-width: 150px;
        object-fit: contain;
    }
    
    /* Estilo para filtros em destaque */
    .filtros-container {
        background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin: 20px 0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .filtros-title {
        font-size: 1.3em;
        font-weight: bold;
        margin-bottom: 15px;
        text-align: center;
    }
    
    .filtros-content {
        background: rgba(255, 255, 255, 0.1);
        padding: 15px;
        border-radius: 8px;
        backdrop-filter: blur(10px);
    }
    
    /* Estilo para métricas */
    .metric-container {
        background: white;
        padding: 20px;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .metric-value {
        font-size: 2em;
        font-weight: bold;
        color: #667eea;
        margin-bottom: 5px;
    }
    
    .metric-label {
        color: #666;
        font-size: 0.9em;
    }
    
    /* Estilo para parágrafos */
    .paragrafo-container {
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        margin: 15px 0;
        overflow: hidden;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .paragrafo-header {
        background: linear-gradient(135deg, #ff9800 0%, #f57c00 100%);
        color: white;
        padding: 12px 20px;
        font-weight: bold;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .paragrafo-content {
        padding: 20px;
        font-family: 'Georgia', 'Times New Roman', serif;
        font-size: 14px;
        line-height: 1.8;
        background: #fafafa;
    }
    
    .paragrafo-numero {
        display: inline-block;
        width: 60px;
        color: #666;
        font-weight: bold;
        font-family: 'Courier New', monospace;
        font-size: 12px;
        margin-right: 15px;
        text-align: right;
    }
    
    .paragrafo-texto {
        display: inline;
    }
    
    .paragrafo-adicionado {
        background-color: #e8f5e8;
        border-left: 4px solid #4caf50;
        color: #2e7d32;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-removido {
        background-color: #ffebee;
        border-left: 4px solid #f44336;
        color: #c62828;
        text-decoration: line-through;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-modificado {
        background-color: #fff3cd;
        border-left: 4px solid #ffc107;
        color: #856404;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-normal {
        background-color: #f9f9f9;
        border-left: 4px solid #e0e0e0;
        color: #555;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .filtro-info {
        background: #e3f2fd;
        border: 1px solid #2196f3;
        border-radius: 6px;
        padding: 10px 15px;
        margin: 10px 0;
        color: #1976d2;
        font-size: 0.9em;
    }
    
    .algoritmo-info {
        background: #f3e5f5;
        border: 1px solid #9c27b0;
        border-radius: 6px;
        padding: 15px;
        margin: 15px 0;
        color: #7b1fa2;
        font-size: 0.9em;
    }
    
    /* Responsividade para mobile */
    @media (max-width: 768px) {
        .header-container {
            flex-direction: column;
            text-align: center;
            gap: 20px;
        }
        
        .header-title {
            font-size: 2em;
        }
        
        .logo-container {
            order: -1;
        }
    }
</style>
""", unsafe_allow_html=True)

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentComparator:
    """Classe principal para comparação de documentos (PDF e Word)"""
    
    def __init__(self):
        self.texto_ref = []
        self.texto_novo = []
        self.diferencas = []
        self.diferencas_detalhadas = []
        self.tipo_ref = None
        self.tipo_novo = None
        
    def detectar_tipo_arquivo(self, nome_arquivo: str) -> str:
        """Detecta o tipo do arquivo baseado na extensão"""
        extensao = Path(nome_arquivo).suffix.lower()
        if extensao == '.pdf':
            return 'pdf'
        elif extensao in ['.docx', '.doc']:
            return 'word'
        else:
            return 'desconhecido'
    
    def validar_arquivo(self, arquivo_bytes: bytes, nome_arquivo: str) -> bool:
        """Valida se o arquivo é válido baseado no tipo"""
        tipo = self.detectar_tipo_arquivo(nome_arquivo)
        
        try:
            if tipo == 'pdf':
                doc = fitz.open(stream=arquivo_bytes, filetype="pdf")
                if doc.page_count == 0:
                    st.error(f"❌ O arquivo PDF '{nome_arquivo}' não contém páginas.")
                    return False
                doc.close()
                return True
                
            elif tipo == 'word':
                if not DOCX_AVAILABLE:
                    st.error("❌ Biblioteca python-docx não está disponível. Instale com: pip install python-docx")
                    return False
                
                # Salvar temporariamente para validar
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(arquivo_bytes)
                    tmp_path = tmp_file.name
                
                try:
                    doc = Document(tmp_path)
                    # Verificar se tem pelo menos um parágrafo
                    if len(doc.paragraphs) == 0:
                        st.error(f"❌ O arquivo Word '{nome_arquivo}' não contém texto.")
                        return False
                    return True
                except Exception as e:
                    st.error(f"❌ Erro ao abrir arquivo Word '{nome_arquivo}': {str(e)}")
                    return False
                finally:
                    # Limpar arquivo temporário
                    try:
                        os.unlink(tmp_path)
                    except:
                        pass
            else:
                st.error(f"❌ Tipo de arquivo não suportado: {nome_arquivo}")
                return False
                
        except Exception as e:
            st.error(f"❌ Erro ao validar '{nome_arquivo}': {str(e)}")
            return False
    
    def extrair_texto_pdf(self, pdf_bytes: bytes) -> List[str]:
        """Extrai texto de cada página do PDF"""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            textos = []
            
            total_paginas = doc.page_count
            
            for i, pagina in enumerate(doc):
                texto = pagina.get_text()
                textos.append(texto)
            
            doc.close()
            return textos
            
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do PDF: {str(e)}")
            return []
    
    def extrair_texto_word(self, word_bytes: bytes) -> List[str]:
        """Extrai texto do documento Word (simula páginas por seções)"""
        try:
            # Salvar temporariamente para processar
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(word_bytes)
                tmp_path = tmp_file.name
            
            try:
                doc = Document(tmp_path)
                
                # Para Word, vamos simular "páginas" agrupando parágrafos
                # Cada "página" terá aproximadamente 50 parágrafos ou quebras de seção
                textos = []
                texto_atual = ""
                contador_paragrafos = 0
                
                for paragrafo in doc.paragraphs:
                    texto_atual += paragrafo.text + "\n"
                    contador_paragrafos += 1
                    
                    # Criar nova "página" a cada 50 parágrafos ou quebra de seção
                    if contador_paragrafos >= 50 or "page-break" in paragrafo.text.lower():
                        if texto_atual.strip():
                            textos.append(texto_atual)
                            texto_atual = ""
                            contador_paragrafos = 0
                
                # Adicionar último texto se houver
                if texto_atual.strip():
                    textos.append(texto_atual)
                
                # Se não há texto, criar pelo menos uma "página" vazia
                if not textos:
                    textos = [""]
                
                return textos
                
            finally:
                # Limpar arquivo temporário
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do Word: {str(e)}")
            return []
    
    def extrair_texto_por_pagina(self, arquivo_bytes: bytes, nome_arquivo: str) -> List[str]:
        """Extrai texto baseado no tipo do arquivo"""
        tipo = self.detectar_tipo_arquivo(nome_arquivo)
        
        progress_bar = st.progress(0)
        
        try:
            if tipo == 'pdf':
                st.info("📖 Extraindo texto do PDF...")
                textos = self.extrair_texto_pdf(arquivo_bytes)
            elif tipo == 'word':
                st.info("📖 Extraindo texto do documento Word...")
                textos = self.extrair_texto_word(arquivo_bytes)
            else:
                st.error(f"❌ Tipo de arquivo não suportado: {tipo}")
                return []
            
            progress_bar.progress(1.0)
            progress_bar.empty()
            return textos
            
        except Exception as e:
            progress_bar.empty()
            st.error(f"❌ Erro ao extrair texto: {str(e)}")
            return []
    
    def normalizar_texto(self, texto: str) -> str:
        """Normaliza o texto removendo variações que não são alterações reais"""
        # Remover espaços extras e quebras de linha desnecessárias
        texto = re.sub(r'\s+', ' ', texto.strip())
        
        # Remover caracteres de controle e formatação
        texto = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', texto)
        
        # Normalizar pontuação (remover espaços antes de pontos, vírgulas, etc.)
        texto = re.sub(r'\s+([,.;:!?])', r'\1', texto)
        
        # Normalizar aspas e caracteres especiais - CORRIGIDO
        texto = re.sub(r'["""]', '"', texto)
        texto = re.sub(r"[''']", "'", texto)
        texto = re.sub(r'[–—]', '-', texto)
        
        return texto
    
    def dividir_em_paragrafos(self, texto: str) -> List[str]:
        """Divide o texto em parágrafos de forma inteligente"""
        # Normalizar o texto primeiro
        texto = self.normalizar_texto(texto)
        
        # Dividir por quebras de linha duplas primeiro (parágrafos naturais)
        paragrafos_brutos = re.split(r'\n\s*\n', texto)
        paragrafos = []
        
        for paragrafo in paragrafos_brutos:
            paragrafo = paragrafo.strip()
            if paragrafo:
                # Se o parágrafo for muito longo (mais de 500 caracteres), dividir por frases
                if len(paragrafo) > 500:
                    # Dividir por pontos finais, mas preservar números decimais
                    frases = re.split(r'(?<!\d)\.(?!\d)\s+', paragrafo)
                    for frase in frases:
                        frase = self.normalizar_texto(frase.strip())
                        if frase and len(frase) > 10:  # Ignorar frases muito curtas
                            paragrafos.append(frase)
                else:
                    paragrafo_normalizado = self.normalizar_texto(paragrafo)
                    if paragrafo_normalizado and len(paragrafo_normalizado) > 10:
                        paragrafos.append(paragrafo_normalizado)
        
        return paragrafos
    
    def calcular_similaridade(self, texto1: str, texto2: str) -> float:
        """Calcula a similaridade entre dois textos (0.0 a 1.0)"""
        if not texto1 and not texto2:
            return 1.0
        if not texto1 or not texto2:
            return 0.0
        
        # Normalizar ambos os textos
        texto1_norm = self.normalizar_texto(texto1)
        texto2_norm = self.normalizar_texto(texto2)
        
        # Usar SequenceMatcher para calcular similaridade
        matcher = difflib.SequenceMatcher(None, texto1_norm, texto2_norm)
        return matcher.ratio()
    
    def encontrar_alteracoes_reais(self, paragrafos_ref: List[str], paragrafos_novo: List[str]) -> List[Dict]:
        """Encontra apenas alterações reais de conteúdo, ignorando deslocamentos"""
        alteracoes = []
        
        # Criar conjuntos de parágrafos únicos para comparação rápida
        set_ref = set(paragrafos_ref)
        set_novo = set(paragrafos_novo)
        
        # Encontrar parágrafos removidos (existem na referência mas não no novo)
        paragrafos_removidos = set_ref - set_novo
        
        # Encontrar parágrafos adicionados (existem no novo mas não na referência)
        paragrafos_adicionados = set_novo - set_ref
        
        # Para parágrafos modificados, precisamos fazer uma análise mais detalhada
        paragrafos_modificados = []
        
        # Verificar se há parágrafos similares que podem ter sido modificados
        for p_ref in paragrafos_removidos.copy():
            melhor_match = None
            melhor_similaridade = 0.0
            
            for p_novo in paragrafos_adicionados.copy():
                similaridade = self.calcular_similaridade(p_ref, p_novo)
                
                # Se a similaridade for alta (>0.6), considerar como modificação
                if similaridade > 0.6 and similaridade > melhor_similaridade:
                    melhor_match = p_novo
                    melhor_similaridade = similaridade
            
            # Se encontrou um match, é uma modificação, não remoção + adição
            if melhor_match and melhor_similaridade > 0.6:
                paragrafos_modificados.append({
                    'original': p_ref,
                    'novo': melhor_match,
                    'similaridade': melhor_similaridade
                })
                paragrafos_removidos.discard(p_ref)
                paragrafos_adicionados.discard(melhor_match)
        
        # Adicionar remoções reais
        for paragrafo in paragrafos_removidos:
            alteracoes.append({
                'tipo': 'removido',
                'texto': paragrafo,
                'texto_original': paragrafo,
                'texto_novo': ''
            })
        
        # Adicionar adições reais
        for paragrafo in paragrafos_adicionados:
            alteracoes.append({
                'tipo': 'adicionado',
                'texto': paragrafo,
                'texto_original': '',
                'texto_novo': paragrafo
            })
        
        # Adicionar modificações reais
        for mod in paragrafos_modificados:
            alteracoes.append({
                'tipo': 'modificado',
                'texto': f"ANTES: {mod['original']}\nDEPOIS: {mod['novo']}",
                'texto_original': mod['original'],
                'texto_novo': mod['novo'],
                'similaridade': mod['similaridade']
            })
        
        return alteracoes
    
    def comparar_textos_por_conteudo(self, texto_ref: List[str], texto_novo: List[str]) -> Tuple[List[Dict], List[Dict]]:
        """Compara textos focando apenas em alterações reais de conteúdo"""
        diferencas_simples = []
        diferencas_detalhadas = []
        
        max_paginas = max(len(texto_ref), len(texto_novo))
        progress_bar = st.progress(0)
        
        for i in range(max_paginas):
            # Garantir que ambos os textos existam
            ref = texto_ref[i] if i < len(texto_ref) else ""
            novo = texto_novo[i] if i < len(texto_novo) else ""
            
            # Dividir em parágrafos
            paragrafos_ref = self.dividir_em_paragrafos(ref)
            paragrafos_novo = self.dividir_em_paragrafos(novo)
            
            # Encontrar alterações reais (não deslocamentos)
            alteracoes = self.encontrar_alteracoes_reais(paragrafos_ref, paragrafos_novo)
            
            if alteracoes:
                # Processar alterações para a tabela simples
                paragrafo_atual = 1
                paragrafos_processados = []
                
                for alteracao in alteracoes:
                    # Adicionar à tabela simples
                    tipo_mapeado = {
                        'removido': 'Removido',
                        'adicionado': 'Adicionado',
                        'modificado': 'Modificado'
                    }[alteracao['tipo']]
                    
                    diferencas_simples.append({
                        'pagina': i + 1,
                        'paragrafo': paragrafo_atual,
                        'tipo': tipo_mapeado,
                        'conteudo_original': alteracao['texto_original'],
                        'conteudo_novo': alteracao['texto_novo']
                    })
                    
                    # Adicionar aos parágrafos processados para visualização
                    paragrafos_processados.append({
                        'numero': paragrafo_atual,
                        'texto': alteracao['texto'],
                        'tipo': alteracao['tipo']
                    })
                    
                    paragrafo_atual += 1
                
                # Adicionar parágrafos inalterados para contexto (limitado)
                paragrafos_inalterados = set(paragrafos_ref) & set(paragrafos_novo)
                contexto_adicionado = 0
                
                for paragrafo in list(paragrafos_inalterados)[:3]:  # Máximo 3 para contexto
                    paragrafos_processados.append({
                        'numero': paragrafo_atual + contexto_adicionado,
                        'texto': paragrafo,
                        'tipo': 'normal'
                    })
                    contexto_adicionado += 1
                
                if paragrafos_processados:
                    diferencas_detalhadas.append({
                        'pagina': i + 1,
                        'paragrafos': paragrafos_processados,
                        'total_paragrafos_ref': len(paragrafos_ref),
                        'total_paragrafos_novo': len(paragrafos_novo),
                        'total_alteracoes': len(alteracoes),
                        'total_contexto': contexto_adicionado
                    })
            
            progress_bar.progress((i + 1) / max_paginas)
        
        progress_bar.empty()
        return diferencas_simples, diferencas_detalhadas

def exibir_diferencas_por_paragrafos(diferencas_detalhadas: List[Dict], tipos_filtro: List[str] = None, paginas_filtro: List[int] = None):
    """Exibe as diferenças por parágrafos com filtros aplicados"""
    if not diferencas_detalhadas:
        st.success("✅ Nenhuma diferença de conteúdo encontrada!")
        st.markdown("""
        <div class="algoritmo-info">
            🎯 <strong>Algoritmo Inteligente:</strong> Este comparador ignora mudanças de posicionamento e formatação, 
            focando apenas em alterações reais de conteúdo. Parágrafos que apenas mudaram de posição não são considerados alterações.
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Aplicar filtros
    diferencas_filtradas = []
    
    for diff_detail in diferencas_detalhadas:
        # Filtrar por página
        if paginas_filtro and diff_detail['pagina'] not in paginas_filtro:
            continue
        
        # Filtrar parágrafos por tipo
        paragrafos_filtrados = []
        for paragrafo in diff_detail['paragrafos']:
            tipo_mapeado = {
                'adicionado': 'Adicionado',
                'removido': 'Removido', 
                'modificado': 'Modificado',
                'normal': 'Normal'
            }.get(paragrafo['tipo'], paragrafo['tipo'])
            
            if not tipos_filtro or any(tipo in tipo_mapeado for tipo in tipos_filtro):
                paragrafos_filtrados.append(paragrafo)
        
        if paragrafos_filtrados:
            diff_filtrada = diff_detail.copy()
            diff_filtrada['paragrafos'] = paragrafos_filtrados
            diff_filtrada['total_alteracoes_filtradas'] = len([p for p in paragrafos_filtrados if p['tipo'] != 'normal'])
            diferencas_filtradas.append(diff_filtrada)
    
    if not diferencas_filtradas:
        st.info("🔍 Nenhuma diferença encontrada com os filtros aplicados.")
        return
    
    st.subheader("🔍 Comparação Visual por Página e Parágrafo")
    
    # Informação sobre o algoritmo
    st.markdown("""
    <div class="algoritmo-info">
        🎯 <strong>Algoritmo Inteligente:</strong> Este comparador ignora mudanças de posicionamento e formatação, 
        focando apenas em alterações reais de conteúdo. Parágrafos similares (>60% de similaridade) são considerados modificações, 
        não remoções + adições separadas.
    </div>
    """, unsafe_allow_html=True)
    
    # Legenda
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("🟢 **Verde:** Conteúdo Adicionado")
    with col2:
        st.markdown("🔴 **Vermelho:** Conteúdo Removido")
    with col3:
        st.markdown("🟡 **Amarelo:** Conteúdo Modificado")
    
    # Informação sobre filtros aplicados
    if tipos_filtro or paginas_filtro:
        filtros_ativos = []
        if tipos_filtro:
            filtros_ativos.append(f"Tipos: {', '.join(tipos_filtro)}")
        if paginas_filtro:
            filtros_ativos.append(f"Páginas: {', '.join(map(str, paginas_filtro))}")
        
        st.markdown(f"""
        <div class="filtro-info">
            🔍 <strong>Filtros aplicados:</strong> {' | '.join(filtros_ativos)}
        </div>
        """, unsafe_allow_html=True)
    
    st.divider()
    
    # Exibir cada página com diferenças
    for diff_detail in diferencas_filtradas:
        st.markdown(f"""
        <div class="paragrafo-container">
            <div class="paragrafo-header">
                <span>🔸 Página/Seção {diff_detail['pagina']}</span>
                <span>{diff_detail.get('total_alteracoes_filtradas', diff_detail['total_alteracoes'])} alteração(ões) de conteúdo | {diff_detail.get('total_contexto', 0)} parágrafo(s) de contexto</span>
            </div>
            <div class="paragrafo-content">
        """, unsafe_allow_html=True)
        
        # Exibir parágrafos
        for paragrafo in diff_detail['paragrafos']:
            tipo_classe = f"paragrafo-{paragrafo['tipo']}"
            
            # Para modificações, mostrar antes e depois
            if paragrafo['tipo'] == 'modificado' and '\nDEPOIS:' in paragrafo['texto']:
                partes = paragrafo['texto'].split('\nDEPOIS:')
                antes = partes[0].replace('ANTES:', '').strip()
                depois = partes[1].strip()
                
                st.markdown(f"""
                    <div class="{tipo_classe}">
                        <span class="paragrafo-numero">§{paragrafo['numero']}</span>
                        <div class="paragrafo-texto">
                            <strong>ANTES:</strong><br>{antes}<br><br>
                            <strong>DEPOIS:</strong><br>{depois}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                    <div class="{tipo_classe}">
                        <span class="paragrafo-numero">§{paragrafo['numero']}</span>
                        <span class="paragrafo-texto">{paragrafo['texto']}</span>
                    </div>
                """, unsafe_allow_html=True)
        
        st.markdown("</div></div>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

def main():
    """Função principal da aplicação"""
    
    # Título e descrição
    st.title("📚 Comparador de PDFs - Solvi")
    st.markdown("**Compare dois documentos (PDF ou Word) e identifique apenas as alterações.**")
    
    # Verificar se python-docx está disponível
    if not DOCX_AVAILABLE:
        st.warning("⚠️ Para suporte completo a documentos Word, instale: `pip install python-docx`")
    
    # Sidebar com informações
    with st.sidebar:
        st.header("ℹ️ Informações")
        st.markdown("""
        **Como usar:**
        1. Faça upload do documento de referência
        2. Faça upload do novo documento
        3. Clique em 'Comparar Documentos'
        4. Visualize apenas as alterações reais de conteúdo
        5. Use os filtros para análise específica
        
        **Formatos suportados:**
        - PDF (.pdf)
        - Word (.docx)
        
        **Algoritmo Inteligente:**
        - ✅ Ignora mudanças de posicionamento
        - ✅ Ignora alterações de formatação
        - ✅ Foca apenas em conteúdo real
        - ✅ Detecta modificações (não remoção + adição)
        - ✅ Normaliza texto para comparação precisa
        
        **Dicas:**
        - Funciona melhor com documentos de texto
        - Parágrafos são normalizados automaticamente
        - Similaridade >60% = modificação, não nova alteração
        """)
    
    # Inicializar o comparador
    if 'comparador' not in st.session_state:
        st.session_state.comparador = DocumentComparator()
    
    # Layout em colunas para upload
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Documento de Referência")
        arquivo_ref = st.file_uploader(
            "Escolha o arquivo de referência",
            type=['pdf', 'docx'] if DOCX_AVAILABLE else ['pdf'],
            key="ref_uploader",
            help="Este será usado como base para comparação"
        )
        
        if arquivo_ref:
            tipo_ref = st.session_state.comparador.detectar_tipo_arquivo(arquivo_ref.name)
            st.success(f"✅ Arquivo carregado: {arquivo_ref.name}")
            st.info(f"📊 Tamanho: {arquivo_ref.size / 1024 / 1024:.2f} MB")
            st.info(f"📋 Tipo: {tipo_ref.upper()}")
    
    with col2:
        st.subheader("📄 Novo Documento")
        arquivo_novo = st.file_uploader(
            "Escolha o novo arquivo",
            type=['pdf', 'docx'] if DOCX_AVAILABLE else ['pdf'],
            key="novo_uploader",
            help="Este será comparado com o arquivo de referência"
        )
        
        if arquivo_novo:
            tipo_novo = st.session_state.comparador.detectar_tipo_arquivo(arquivo_novo.name)
            st.success(f"✅ Arquivo carregado: {arquivo_novo.name}")
            st.info(f"📊 Tamanho: {arquivo_novo.size / 1024 / 1024:.2f} MB")
            st.info(f"📋 Tipo: {tipo_novo.upper()}")
    
    # Botão de comparação
    if arquivo_ref and arquivo_novo:
        # Verificar se os tipos são compatíveis
        tipo_ref = st.session_state.comparador.detectar_tipo_arquivo(arquivo_ref.name)
        tipo_novo = st.session_state.comparador.detectar_tipo_arquivo(arquivo_novo.name)
        
        if tipo_ref != tipo_novo:
            st.warning(f"⚠️ Atenção: Você está comparando arquivos de tipos diferentes ({tipo_ref.upper()} vs {tipo_novo.upper()}). A comparação ainda é possível, mas pode não ser ideal.")
        
        if st.button("🔍 Comparar Documentos", type="primary", use_container_width=True):
            
            with st.spinner("🔄 Processando arquivos..."):
                # Validar arquivos
                ref_bytes = arquivo_ref.read()
                novo_bytes = arquivo_novo.read()
                
                if not st.session_state.comparador.validar_arquivo(ref_bytes, arquivo_ref.name):
                    st.stop()
                
                if not st.session_state.comparador.validar_arquivo(novo_bytes, arquivo_novo.name):
                    st.stop()
                
                # Extrair textos
                texto_ref = st.session_state.comparador.extrair_texto_por_pagina(ref_bytes, arquivo_ref.name)
                texto_novo = st.session_state.comparador.extrair_texto_por_pagina(novo_bytes, arquivo_novo.name)
                
                if not texto_ref or not texto_novo:
                    st.error("❌ Erro ao extrair texto dos documentos")
                    st.stop()
                
                # Comparar textos focando apenas em conteúdo
                st.info("🔍 Analisando alterações reais de conteúdo...")
                diferencas_simples, diferencas_detalhadas = st.session_state.comparador.comparar_textos_por_conteudo(texto_ref, texto_novo)
                
                # Armazenar resultados no session state
                st.session_state.diferencas = diferencas_simples
                st.session_state.diferencas_detalhadas = diferencas_detalhadas
                st.session_state.arquivo_ref_nome = arquivo_ref.name
                st.session_state.arquivo_novo_nome = arquivo_novo.name
                st.session_state.tipo_ref = tipo_ref
                st.session_state.tipo_novo = tipo_novo
    
    # Exibir resultados se existirem
    if 'diferencas' in st.session_state and 'diferencas_detalhadas' in st.session_state:
        diferencas = st.session_state.diferencas
        diferencas_detalhadas = st.session_state.diferencas_detalhadas
        
        st.divider()
        
        # Resumo dos resultados com layout melhorado
        st.subheader("📊 Resumo da Análise")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{len(diferencas)}</div>
                <div class="metric-label">Alterações Reais</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            paginas_afetadas = len(set(d['pagina'] for d in diferencas)) if diferencas else 0
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{paginas_afetadas}</div>
                <div class="metric-label">Páginas/Seções Afetadas</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            tipos_mudanca = len(set(d['tipo'] for d in diferencas)) if diferencas else 0
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{tipos_mudanca}</div>
                <div class="metric-label">Tipos de Mudança</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            compatibilidade = "✅ Mesmos tipos" if st.session_state.tipo_ref == st.session_state.tipo_novo else "⚠️ Tipos diferentes"
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value" style="font-size: 1.2em;">{compatibilidade}</div>
                <div class="metric-label">Compatibilidade</div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        if diferencas:
            # Filtros Avançados em destaque
            st.markdown("""
            <div class="filtros-container">
                <div class="filtros-title">🔍 Filtros Avançados</div>
                <div class="filtros-content">
                    <p style="margin-bottom: 15px; text-align: center;">Use os filtros abaixo para analisar tipos específicos de alterações reais de conteúdo</p>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Converter para DataFrame para melhor visualização
            df_diferencas = pd.DataFrame(diferencas)
            
            col1, col2 = st.columns(2)
            
            with col1:
                tipos_selecionados = st.multiselect(
                    "🏷️ Filtrar por tipo de mudança:",
                    options=df_diferencas['tipo'].unique(),
                    default=df_diferencas['tipo'].unique(),
                    help="Selecione os tipos de alteração que deseja visualizar"
                )
            
            with col2:
                paginas_selecionadas = st.multiselect(
                    "📄 Filtrar por página/seção:",
                    options=sorted(df_diferencas['pagina'].unique()),
                    default=sorted(df_diferencas['pagina'].unique()),
                    help="Selecione as páginas/seções que deseja analisar"
                )
            
            # Aplicar filtros
            df_filtrado = df_diferencas[
                (df_diferencas['tipo'].isin(tipos_selecionados)) &
                (df_diferencas['pagina'].isin(paginas_selecionadas))
            ]
            
            # Exibir comparação visual com filtros aplicados
            exibir_diferencas_por_paragrafos(diferencas_detalhadas, tipos_selecionados, paginas_selecionadas)
            
            # Tabela Resumo das Diferenças (retrátil)
            with st.expander("📋 Tabela Resumo das Diferenças", expanded=False):
                if len(df_filtrado) != len(df_diferencas):
                    st.info(f"📊 Mostrando {len(df_filtrado)} de {len(df_diferencas)} diferenças (filtros aplicados)")
                
                # Configurar exibição da tabela
                st.dataframe(
                    df_filtrado,
                    use_container_width=True,
                    column_config={
                        "pagina": st.column_config.NumberColumn("Página/Seção", format="%d"),
                        "paragrafo": st.column_config.NumberColumn("Parágrafo", format="%d"),
                        "tipo": st.column_config.TextColumn("Tipo"),
                        "conteudo_original": st.column_config.TextColumn("Conteúdo Original"),
                        "conteudo_novo": st.column_config.TextColumn("Conteúdo Novo")
                    }
                )
                
                # Estatísticas dos dados filtrados
                if len(df_filtrado) > 0:
                    st.markdown("### 📈 Estatísticas dos Dados Filtrados")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Total de Alterações", len(df_filtrado))
                    
                    with col2:
                        paginas_filtradas = len(df_filtrado['pagina'].unique())
                        st.metric("Páginas Afetadas", paginas_filtradas)
                    
                    with col3:
                        tipos_filtrados = len(df_filtrado['tipo'].unique())
                        st.metric("Tipos de Mudança", tipos_filtrados)
        else:
            # Exibir comparação visual mesmo sem diferenças
            exibir_diferencas_por_paragrafos(diferencas_detalhadas)
        
        if not diferencas:
            st.balloons()

if __name__ == "__main__":
    main()

