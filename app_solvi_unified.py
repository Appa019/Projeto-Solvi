"""
🌱 Plataforma Solví - Análise Inteligente de Documentos
Versão com Visualização Avançada de Diferenças -  Visual Diff
Aplicação  que combina análise CVM e comparação visual de documentos
"""

import streamlit as st
import pandas as pd
import openai
from io import BytesIO
import PyPDF2
import docx
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import json
import time
import re
from datetime import datetime
import base64
import fitz  # PyMuPDF
import difflib
from typing import List, Tuple, Dict, Optional, Set
import logging
from pathlib import Path
import tempfile
import os
import html

# Configuração da página com tema Solví - SIDEBAR SEMPRE EXPANDIDA
st.set_page_config(
    page_title="Plataforma Solví - Soluções Inteligentes",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://www.solvi.com',
        'Report a bug': 'https://www.solvi.com/contato',
        'About': "Plataforma Solví - Soluções para a vida com sustentabilidade e inovação"
    }
)

# CSS  Masterpiece + VISUALIZAÇÃO AVANÇADA DE DIFERENÇAS
st.markdown("""
<style>
    /* Importar fontes oficiais */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700;800;900&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@300;400;500;600;700;800&display=swap');
    
    /* Reset completo e configurações globais */
    * {
        box-sizing: border-box;
        margin: 0;
        padding: 0;
    }
    
    /* CORREÇÃO PARA HEADER PREENCHER TOPO COMPLETO */
    .stApp {
        margin-top: 0px !important;
        padding-top: 0px !important;
    }
    
    .stApp > header {
        display: none !important;
    }
    
    .main .block-container {
        padding-top: 0rem !important;
        padding-bottom: 3rem;
        max-width: 1400px;
        padding-left: 2rem;
        padding-right: 2rem;
        margin-top: 0rem !important;
    }
    
    /* CORREÇÃO PARA SIDEBAR SEMPRE VISÍVEL */
    section[data-testid="stSidebar"] {
        width: 21rem !important;
        min-width: 21rem !important;
        display: block !important;
        visibility: visible !important;
        background: var(--solvi-gradient-surface) !important;
        border-right: 3px solid var(--solvi-light-green) !important;
    }
    
    section[data-testid="stSidebar"] > div {
        width: 21rem !important;
        min-width: 21rem !important;
        background: var(--solvi-gradient-surface) !important;
        padding: 2rem 1.5rem !important;
    }
    
    .css-1d391kg {
        width: 21rem !important;
        min-width: 21rem !important;
        background: var(--solvi-gradient-surface) !important;
    }
    
    /* Ocultar botão de toggle da sidebar */
    button[kind="header"] {
        display: none !important;
    }
    
    /* Ocultar elementos padrão do Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {visibility: hidden;}
    
    /* Paleta de cores oficial Solví - Verde Escuro Dominante */
    :root {
        --solvi-dark-green: #0d4f1c;
        --solvi-primary-green: #1b5e20;
        --solvi-medium-green: #2e7d32;
        --solvi-light-green: #388e3c;
        --solvi-accent-green: #4caf50;
        --solvi-bright-green: #66bb6a;
        --solvi-surface: #f1f8e9;
        --solvi-background: #e8f5e8;
        --solvi-white: #ffffff;
        --solvi-text-dark: #1b5e20;
        --solvi-text-light: #ffffff;
        --solvi-shadow: rgba(13, 79, 28, 0.15);
        --solvi-shadow-strong: rgba(13, 79, 28, 0.25);
        --solvi-gradient-primary: linear-gradient(135deg, var(--solvi-dark-green) 0%, var(--solvi-primary-green) 30%, var(--solvi-medium-green) 70%, var(--solvi-light-green) 100%);
        --solvi-gradient-surface: linear-gradient(135deg, var(--solvi-white) 0%, var(--solvi-surface) 50%, var(--solvi-background) 100%);
        
        /* Cores para visualização de diferenças */
        --diff-added: #d4edda;
        --diff-added-border: #28a745;
        --diff-added-text: #155724;
        --diff-removed: #f8d7da;
        --diff-removed-border: #dc3545;
        --diff-removed-text: #721c24;
        --diff-modified: #fff3cd;
        --diff-modified-border: #ffc107;
        --diff-modified-text: #856404;
        --diff-unchanged: #f8f9fa;
        --diff-unchanged-border: #dee2e6;
        --diff-unchanged-text: #495057;
    }
    
    /* Header Masterpiece - Largura total e impacto visual máximo */
    .solvi-header {
        background: var(--solvi-gradient-primary);
        color: var(--solvi-text-light);
        padding: 3rem 0;
        border-radius: 0;
        margin: -3rem calc(-50vw + 50%) 3rem calc(-50vw + 50%);
        margin-top: -3rem !important;
        box-shadow: 0 12px 40px var(--solvi-shadow-strong);
        position: relative;
        overflow: hidden;
        min-height: 200px;
        width: 100vw;
        z-index: 10;
    }
    
    .solvi-header::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: url('https://images.unsplash.com/photo-1441974231531-c6227db76b6e?ixlib=rb-4.0.3&auto=format&fit=crop&w=1920&q=80');
        background-size: cover;
        background-position: center;
        opacity: 0.08;
        z-index: 0;
    }
    
    .solvi-header::after {
        content: '';
        position: absolute;
        bottom: 0;
        left: 0;
        right: 0;
        height: 8px;
        background: linear-gradient(90deg, 
            var(--solvi-bright-green) 0%, 
            var(--solvi-accent-green) 25%, 
            var(--solvi-light-green) 50%, 
            var(--solvi-medium-green) 75%, 
            var(--solvi-primary-green) 100%);
        z-index: 1;
    }
    
    .solvi-header-content {
        position: relative;
        z-index: 2;
        display: flex;
        align-items: center;
        justify-content: space-between;
        flex-wrap: wrap;
        gap: 2.5rem;
        max-width: 1400px;
        margin: 0 auto;
        padding: 0 3rem;
    }
    
    .solvi-logo-section {
        display: flex;
        align-items: center;
        gap: 2.5rem;
        flex: 1;
    }
    
    /* Logo  com Background Verde Escuro */
    .solvi-logo {
        height: 80px;
        width: auto;
        background: var(--solvi-dark-green);
        padding: 15px 25px;
        border-radius: 16px;
        box-shadow: 0 8px 30px rgba(0,0,0,0.2);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        border: 4px solid var(--solvi-primary-green);
        filter: brightness(1.1);
    }
    
    .solvi-logo:hover {
        transform: scale(1.08) rotate(2deg);
        box-shadow: 0 12px 40px rgba(0,0,0,0.3);
        background: var(--solvi-primary-green);
        border-color: var(--solvi-medium-green);
        filter: brightness(1.2);
    }
    
    .solvi-title-section {
        flex: 2;
    }
    
    .solvi-title {
        font-size: 3.5rem;
        font-weight: 900;
        font-family: 'Poppins', sans-serif;
        margin: 0;
        text-shadow: 3px 3px 12px rgba(0,0,0,0.4);
        letter-spacing: -2px;
        line-height: 1;
        background: linear-gradient(45deg, #ffffff 0%, #f0f8ff 50%, #ffffff 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .solvi-subtitle {
        font-size: 1.4rem;
        opacity: 0.95;
        margin-top: 0.8rem;
        font-weight: 600;
        letter-spacing: 0.8px;
        text-shadow: 2px 2px 6px rgba(0,0,0,0.3);
        font-family: 'Inter', sans-serif;
    }
    
    .solvi-badge {
        background: rgba(255,255,255,0.15);
        padding: 1.5rem 3rem;
        border-radius: 50px;
        font-size: 1.2rem;
        font-weight: 800;
        backdrop-filter: blur(20px);
        border: 3px solid rgba(255,255,255,0.25);
        text-transform: uppercase;
        letter-spacing: 2px;
        box-shadow: 0 6px 25px rgba(0,0,0,0.15);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        font-family: 'Poppins', sans-serif;
        white-space: nowrap;
    }
    
    .solvi-badge:hover {
        background: rgba(255,255,255,0.25);
        transform: translateY(-3px) scale(1.05);
        box-shadow: 0 10px 35px rgba(0,0,0,0.2);
        border-color: rgba(255,255,255,0.4);
    }
    
    /* VISUALIZAÇÃO AVANÇADA DE DIFERENÇAS -  */
    .diff-viewer {
        background: var(--solvi-white);
        border-radius: 20px;
        border: 3px solid var(--solvi-background);
        box-shadow: 0 15px 50px var(--solvi-shadow);
        margin: 3rem 0;
        overflow: hidden;
        position: relative;
    }
    
    .diff-viewer::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 6px;
        background: var(--solvi-gradient-primary);
        z-index: 1;
    }
    
    .diff-header {
        background: var(--solvi-gradient-surface);
        padding: 2rem 3rem;
        border-bottom: 3px solid var(--solvi-background);
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: wrap;
        gap: 2rem;
    }
    
    .diff-title {
        font-size: 1.8rem;
        font-weight: 800;
        color: var(--solvi-text-dark);
        font-family: 'Poppins', sans-serif;
        margin: 0;
    }
    
    .diff-stats {
        display: flex;
        gap: 2rem;
        flex-wrap: wrap;
    }
    
    .diff-stat {
        display: flex;
        align-items: center;
        gap: 0.8rem;
        padding: 0.8rem 1.5rem;
        border-radius: 12px;
        font-weight: 700;
        font-family: 'Inter', sans-serif;
        font-size: 1rem;
        border: 2px solid;
        transition: all 0.3s ease;
    }
    
    .diff-stat:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
    }
    
    .diff-stat.added {
        background: var(--diff-added);
        border-color: var(--diff-added-border);
        color: var(--diff-added-text);
    }
    
    .diff-stat.removed {
        background: var(--diff-removed);
        border-color: var(--diff-removed-border);
        color: var(--diff-removed-text);
    }
    
    .diff-stat.modified {
        background: var(--diff-modified);
        border-color: var(--diff-modified-border);
        color: var(--diff-modified-text);
    }
    
    .diff-content {
        max-height: 600px;
        overflow-y: auto;
        padding: 0;
    }
    
    .diff-line {
        display: flex;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.95rem;
        line-height: 1.6;
        border-bottom: 1px solid #f0f0f0;
        transition: all 0.2s ease;
        position: relative;
    }
    
    .diff-line:hover {
        background: rgba(0,0,0,0.02);
    }
    
    .diff-line-number {
        width: 80px;
        padding: 0.8rem 1rem;
        background: #f8f9fa;
        border-right: 2px solid #dee2e6;
        text-align: center;
        font-weight: 600;
        color: #6c757d;
        user-select: none;
        flex-shrink: 0;
    }
    
    .diff-line-content {
        flex: 1;
        padding: 0.8rem 1.5rem;
        white-space: pre-wrap;
        word-wrap: break-word;
        position: relative;
    }
    
    /* Tipos de diferenças */
    .diff-line.added {
        background: var(--diff-added);
        border-left: 4px solid var(--diff-added-border);
    }
    
    .diff-line.added .diff-line-number {
        background: var(--diff-added);
        color: var(--diff-added-text);
        font-weight: 800;
    }
    
    .diff-line.added .diff-line-content {
        color: var(--diff-added-text);
        font-weight: 600;
    }
    
    .diff-line.added::before {
        content: '+';
        position: absolute;
        left: 85px;
        top: 50%;
        transform: translateY(-50%);
        color: var(--diff-added-border);
        font-weight: 900;
        font-size: 1.2rem;
        z-index: 1;
    }
    
    .diff-line.removed {
        background: var(--diff-removed);
        border-left: 4px solid var(--diff-removed-border);
    }
    
    .diff-line.removed .diff-line-number {
        background: var(--diff-removed);
        color: var(--diff-removed-text);
        font-weight: 800;
    }
    
    .diff-line.removed .diff-line-content {
        color: var(--diff-removed-text);
        font-weight: 600;
        text-decoration: line-through;
        opacity: 0.8;
    }
    
    .diff-line.removed::before {
        content: '-';
        position: absolute;
        left: 85px;
        top: 50%;
        transform: translateY(-50%);
        color: var(--diff-removed-border);
        font-weight: 900;
        font-size: 1.2rem;
        z-index: 1;
    }
    
    .diff-line.modified {
        background: var(--diff-modified);
        border-left: 4px solid var(--diff-modified-border);
    }
    
    .diff-line.modified .diff-line-number {
        background: var(--diff-modified);
        color: var(--diff-modified-text);
        font-weight: 800;
    }
    
    .diff-line.modified .diff-line-content {
        color: var(--diff-modified-text);
        font-weight: 600;
    }
    
    .diff-line.modified::before {
        content: '~';
        position: absolute;
        left: 85px;
        top: 50%;
        transform: translateY(-50%);
        color: var(--diff-modified-border);
        font-weight: 900;
        font-size: 1.2rem;
        z-index: 1;
    }
    
    .diff-line.unchanged {
        background: var(--diff-unchanged);
    }
    
    .diff-line.unchanged .diff-line-number {
        background: var(--diff-unchanged);
        color: var(--diff-unchanged-text);
    }
    
    .diff-line.unchanged .diff-line-content {
        color: var(--diff-unchanged-text);
        opacity: 0.7;
    }
    
    /* Highlights dentro do texto */
    .diff-highlight {
        padding: 0.2rem 0.4rem;
        border-radius: 4px;
        font-weight: 700;
        position: relative;
    }
    
    .diff-highlight.added {
        background: #28a745;
        color: white;
    }
    
    .diff-highlight.removed {
        background: #dc3545;
        color: white;
        text-decoration: line-through;
    }
    
    .diff-highlight.modified {
        background: #ffc107;
        color: #856404;
    }
    
    /* Seção de resumo visual */
    .diff-summary {
        background: var(--solvi-gradient-surface);
        padding: 3rem;
        border-top: 3px solid var(--solvi-background);
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 2rem;
    }
    
    .diff-summary-item {
        background: var(--solvi-white);
        padding: 2rem;
        border-radius: 16px;
        border: 2px solid var(--solvi-background);
        text-align: center;
        transition: all 0.3s ease;
        box-shadow: 0 5px 15px rgba(0,0,0,0.05);
    }
    
    .diff-summary-item:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
        border-color: var(--solvi-light-green);
    }
    
    .diff-summary-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
        display: block;
    }
    
    .diff-summary-value {
        font-size: 2.5rem;
        font-weight: 900;
        font-family: 'Poppins', sans-serif;
        margin-bottom: 0.5rem;
        background: var(--solvi-gradient-primary);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .diff-summary-label {
        font-size: 1.1rem;
        font-weight: 600;
        color: #666;
        font-family: 'Inter', sans-serif;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    /* Controles de visualização */
    .diff-controls {
        background: var(--solvi-gradient-surface);
        padding: 2rem 3rem;
        border-bottom: 3px solid var(--solvi-background);
        display: flex;
        gap: 1.5rem;
        flex-wrap: wrap;
        align-items: center;
    }
    
    .diff-control-group {
        display: flex;
        align-items: center;
        gap: 1rem;
    }
    
    .diff-control-label {
        font-weight: 600;
        color: var(--solvi-text-dark);
        font-family: 'Inter', sans-serif;
        font-size: 0.95rem;
    }
    
    .diff-toggle {
        display: flex;
        background: var(--solvi-white);
        border-radius: 8px;
        border: 2px solid var(--solvi-background);
        overflow: hidden;
    }
    
    .diff-toggle-btn {
        padding: 0.6rem 1.2rem;
        border: none;
        background: transparent;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
        font-size: 0.9rem;
        cursor: pointer;
        transition: all 0.2s ease;
        color: #666;
    }
    
    .diff-toggle-btn.active {
        background: var(--solvi-primary-green);
        color: white;
    }
    
    .diff-toggle-btn:hover:not(.active) {
        background: var(--solvi-surface);
    }
    
    /* Seção de imagens inspiracionais  */
    .solvi-inspiration {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 2.5rem;
        margin: 4rem 0;
        padding: 3rem;
        background: var(--solvi-gradient-surface);
        border-radius: 28px;
        border: 3px solid var(--solvi-light-green);
        box-shadow: 0 12px 40px var(--solvi-shadow);
        position: relative;
        overflow: hidden;
    }
    
    .solvi-inspiration::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 6px;
        background: var(--solvi-gradient-primary);
        z-index: 1;
    }
    
    .solvi-inspiration-item {
        position: relative;
        border-radius: 20px;
        overflow: hidden;
        box-shadow: 0 10px 35px var(--solvi-shadow);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        background: var(--solvi-white);
        border: 3px solid var(--solvi-background);
        transform-origin: center;
    }
    
    .solvi-inspiration-item:hover {
        transform: translateY(-12px) scale(1.02);
        box-shadow: 0 20px 60px var(--solvi-shadow-strong);
        border-color: var(--solvi-light-green);
    }
    
    .solvi-inspiration-image {
        width: 100%;
        height: 200px;
        object-fit: cover;
        border-radius: 20px 20px 0 0;
        transition: all 0.4s ease;
    }
    
    .solvi-inspiration-item:hover .solvi-inspiration-image {
        transform: scale(1.05);
        filter: brightness(1.1) saturate(1.2);
    }
    
    .solvi-inspiration-content {
        padding: 2.5rem;
        background: var(--solvi-gradient-surface);
        position: relative;
    }
    
    .solvi-inspiration-title {
        font-size: 1.4rem;
        font-weight: 800;
        color: var(--solvi-text-dark);
        margin: 0 0 1.2rem 0;
        font-family: 'Poppins', sans-serif;
        letter-spacing: -0.5px;
    }
    
    .solvi-inspiration-desc {
        font-size: 1.05rem;
        color: #555;
        line-height: 1.7;
        margin: 0;
        font-weight: 500;
        font-family: 'Inter', sans-serif;
    }
    
    /* Navegação  estilo Solví */
    .solvi-navigation {
        background: var(--solvi-gradient-surface);
        border-radius: 24px;
        padding: 2rem;
        margin: 4rem 0;
        box-shadow: 0 12px 45px var(--solvi-shadow);
        border: 3px solid var(--solvi-background);
        position: relative;
        overflow: hidden;
    }
    
    .solvi-navigation::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 6px;
        background: var(--solvi-gradient-primary);
        z-index: 1;
    }
    
    /* Cards  estilo Solví */
    .solvi-card {
        background: var(--solvi-gradient-surface);
        border-radius: 28px;
        padding: 3.5rem;
        margin: 3rem 0;
        box-shadow: 0 20px 60px var(--solvi-shadow);
        border: 3px solid var(--solvi-background);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        overflow: hidden;
    }
    
    .solvi-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 8px;
        background: var(--solvi-gradient-primary);
        z-index: 1;
    }
    
    .solvi-card:hover {
        transform: translateY(-10px);
        box-shadow: 0 30px 80px var(--solvi-shadow-strong);
        border-color: var(--solvi-light-green);
    }
    
    .solvi-card-header {
        display: flex;
        align-items: center;
        margin-bottom: 3rem;
        padding-bottom: 2.5rem;
        border-bottom: 4px solid var(--solvi-background);
        position: relative;
    }
    
    .solvi-card-icon {
        width: 80px;
        height: 80px;
        background: var(--solvi-gradient-primary);
        border-radius: 24px;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-right: 2.5rem;
        font-size: 2.5rem;
        box-shadow: 0 10px 30px var(--solvi-shadow);
        border: 4px solid var(--solvi-white);
        transition: all 0.3s ease;
    }
    
    .solvi-card-icon:hover {
        transform: scale(1.1) rotate(5deg);
        box-shadow: 0 15px 40px var(--solvi-shadow-strong);
    }
    
    .solvi-card-title {
        font-size: 2.2rem;
        font-weight: 900;
        color: var(--solvi-text-dark);
        margin: 0;
        font-family: 'Poppins', sans-serif;
        letter-spacing: -1px;
        line-height: 1.2;
    }
    
    /* Métricas  estilo Solví */
    .solvi-metrics {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
        gap: 3rem;
        margin: 4rem 0;
    }
    
    .solvi-metric {
        background: var(--solvi-gradient-surface);
        border-radius: 28px;
        padding: 3.5rem;
        text-align: center;
        border: 3px solid var(--solvi-background);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        overflow: hidden;
        box-shadow: 0 15px 50px var(--solvi-shadow);
    }
    
    .solvi-metric::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 8px;
        background: var(--solvi-gradient-primary);
        z-index: 1;
    }
    
    .solvi-metric:hover {
        transform: translateY(-12px) scale(1.02);
        box-shadow: 0 25px 70px var(--solvi-shadow-strong);
        border-color: var(--solvi-light-green);
    }
    
    .solvi-metric-value {
        font-size: 4.5rem;
        font-weight: 900;
        color: var(--solvi-primary-green);
        margin-bottom: 1.5rem;
        font-family: 'Poppins', sans-serif;
        line-height: 1;
        text-shadow: 3px 3px 6px rgba(0,0,0,0.1);
        background: var(--solvi-gradient-primary);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .solvi-metric-label {
        color: #555;
        font-size: 1.2rem;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 2px;
        font-family: 'Inter', sans-serif;
    }
    
    /* Botões  estilo Solví */
    .stButton > button {
        background: var(--solvi-gradient-primary);
        color: var(--solvi-text-light);
        border: none;
        border-radius: 24px;
        padding: 1.8rem 4rem;
        font-weight: 800;
        font-size: 1.3rem;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 10px 30px var(--solvi-shadow);
        font-family: 'Poppins', sans-serif;
        letter-spacing: 1.5px;
        text-transform: uppercase;
        border: 4px solid transparent;
        position: relative;
        overflow: hidden;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
        transition: left 0.5s;
    }
    
    .stButton > button:hover::before {
        left: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-6px) scale(1.02);
        box-shadow: 0 15px 50px var(--solvi-shadow-strong);
        background: linear-gradient(135deg, var(--solvi-dark-green) 0%, var(--solvi-primary-green) 50%, var(--solvi-medium-green) 100%);
        border-color: var(--solvi-white);
    }
    
    /* Alertas  estilo Solví */
    .solvi-alert {
        border-radius: 24px;
        padding: 3rem 3.5rem;
        margin: 3rem 0;
        border-left: 10px solid;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
        position: relative;
        overflow: hidden;
        box-shadow: 0 12px 40px rgba(0,0,0,0.1);
        font-size: 1.15rem;
        line-height: 1.6;
    }
    
    .solvi-alert.success {
        background: var(--solvi-gradient-surface);
        border-color: var(--solvi-accent-green);
        color: var(--solvi-text-dark);
    }
    
    .solvi-alert.warning {
        background: linear-gradient(135deg, #fff8e1 0%, #fffde7 50%, #f9fbe7 100%);
        border-color: #ff9800;
        color: #e65100;
    }
    
    .solvi-alert.error {
        background: linear-gradient(135deg, #ffebee 0%, #fce4ec 50%, #f3e5f5 100%);
        border-color: #f44336;
        color: #c62828;
    }
    
    .solvi-alert.info {
        background: linear-gradient(135deg, #e3f2fd 0%, #e1f5fe 50%, #e0f2f1 100%);
        border-color: #2196f3;
        color: #0d47a1;
    }
    
    /* Upload areas  com pontilhado contínuo */
    .solvi-upload {
        border: 5px dashed var(--solvi-light-green);
        border-radius: 28px;
        padding: 5rem 4rem;
        text-align: center;
        background: var(--solvi-gradient-surface);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        margin: 3rem 0;
        box-shadow: 0 15px 50px var(--solvi-shadow);
        position: relative;
        overflow: hidden;
    }
    
    /* Pontilhado contínuo animado  */
    .solvi-upload::before {
        content: '';
        position: absolute;
        top: -3px;
        left: -3px;
        right: -3px;
        bottom: -3px;
        border: 5px dashed var(--solvi-primary-green);
        border-radius: 28px;
        animation: dash 25s linear infinite;
        opacity: 0.7;
        z-index: 0;
    }
    
    @keyframes dash {
        0% {
            stroke-dashoffset: 0;
            transform: rotate(0deg);
        }
        100% {
            stroke-dashoffset: 50px;
            transform: rotate(360deg);
        }
    }
    
    .solvi-upload:hover {
        border-color: var(--solvi-primary-green);
        background: linear-gradient(135deg, var(--solvi-surface) 0%, var(--solvi-background) 50%, #dcedc8 100%);
        transform: translateY(-8px) scale(1.01);
        box-shadow: 0 20px 70px var(--solvi-shadow-strong);
    }
    
    .solvi-upload:hover::before {
        border-color: var(--solvi-dark-green);
        opacity: 0.9;
        animation-duration: 15s;
    }
    
    .solvi-upload-icon {
        font-size: 5rem;
        color: var(--solvi-light-green);
        margin-bottom: 2.5rem;
        text-shadow: 3px 3px 6px rgba(0,0,0,0.1);
        position: relative;
        z-index: 1;
        transition: all 0.3s ease;
    }
    
    .solvi-upload:hover .solvi-upload-icon {
        transform: scale(1.1) rotate(5deg);
        color: var(--solvi-primary-green);
    }
    
    .solvi-upload-text {
        font-size: 1.8rem;
        font-weight: 800;
        color: var(--solvi-text-dark);
        margin-bottom: 1.5rem;
        font-family: 'Poppins', sans-serif;
        position: relative;
        z-index: 1;
        letter-spacing: -0.5px;
    }
    
    .solvi-upload-subtext {
        font-size: 1.2rem;
        color: #666;
        line-height: 1.8;
        font-weight: 500;
        position: relative;
        z-index: 1;
        font-family: 'Inter', sans-serif;
    }
    
    /* Footer  */
    .solvi-footer {
        background: var(--solvi-gradient-primary);
        color: var(--solvi-text-light);
        padding: 5rem 3rem;
        border-radius: 28px;
        margin: 5rem 0 3rem 0;
        text-align: center;
        box-shadow: 0 20px 60px var(--solvi-shadow-strong);
        position: relative;
        overflow: hidden;
    }
    
    .solvi-footer::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: url('https://images.unsplash.com/photo-1542601906990-b4d3fb778b09?ixlib=rb-4.0.3&auto=format&fit=crop&w=1920&q=80');
        background-size: cover;
        background-position: center;
        opacity: 0.06;
        z-index: 0;
    }
    
    .solvi-footer-content {
        position: relative;
        z-index: 1;
    }
    
    .solvi-footer-logo {
        height: 70px;
        margin-bottom: 2.5rem;
        background: var(--solvi-dark-green);
        padding: 15px 25px;
        border-radius: 16px;
        box-shadow: 0 8px 25px rgba(0,0,0,0.3);
        border: 4px solid var(--solvi-primary-green);
        transition: all 0.3s ease;
    }
    
    .solvi-footer-logo:hover {
        transform: scale(1.05);
        box-shadow: 0 12px 35px rgba(0,0,0,0.4);
    }
    
    /* SIDEBAR  STYLING */
    .sidebar-content {
        background: var(--solvi-gradient-surface);
        border-radius: 20px;
        padding: 2rem;
        margin-bottom: 2rem;
        border: 3px solid var(--solvi-background);
        box-shadow: 0 10px 30px var(--solvi-shadow);
    }
    
    .sidebar-title {
        color: var(--solvi-primary-green);
        font-family: 'Poppins', sans-serif;
        font-weight: 800;
        font-size: 1.4rem;
        margin-bottom: 1.5rem;
        text-align: center;
        border-bottom: 3px solid var(--solvi-light-green);
        padding-bottom: 1rem;
    }
    
    /* Responsividade  */
    @media (max-width: 768px) {
        .solvi-header {
            padding: 2.5rem 0;
            min-height: 180px;
        }
        
        .solvi-header-content {
            flex-direction: column;
            text-align: center;
            gap: 2rem;
            padding: 0 1.5rem;
        }
        
        .solvi-title {
            font-size: 2.5rem;
        }
        
        .solvi-subtitle {
            font-size: 1.2rem;
        }
        
        .solvi-logo {
            height: 70px;
        }
        
        .solvi-inspiration {
            grid-template-columns: 1fr;
            padding: 2.5rem;
        }
        
        .solvi-metrics {
            grid-template-columns: 1fr;
        }
        
        .solvi-metric-value {
            font-size: 4rem;
        }
        
        .solvi-card {
            padding: 3rem;
        }
        
        .solvi-upload {
            padding: 4rem 2.5rem;
        }
        
        .solvi-badge {
            padding: 1.2rem 2.5rem;
            font-size: 1rem;
        }
        
        /* Sidebar responsiva */
        section[data-testid="stSidebar"] {
            width: 18rem !important;
            min-width: 18rem !important;
        }
        
        section[data-testid="stSidebar"] > div {
            width: 18rem !important;
            min-width: 18rem !important;
            padding: 1.5rem 1rem !important;
        }
        
        /* Diff viewer responsivo */
        .diff-line-number {
            width: 60px;
            font-size: 0.8rem;
        }
        
        .diff-line-content {
            font-size: 0.85rem;
            padding: 0.6rem 1rem;
        }
        
        .diff-stats {
            flex-direction: column;
            gap: 1rem;
        }
        
        .diff-controls {
            flex-direction: column;
            gap: 1rem;
        }
    }
    
    /* Animações  */
    @keyframes fadeInUp {
        from {
            opacity: 0;
            transform: translateY(60px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    @keyframes pulse {
        0%, 100% {
            transform: scale(1);
        }
        50% {
            transform: scale(1.03);
        }
    }
    
    @keyframes float {
        0%, 100% {
            transform: translateY(0px);
        }
        50% {
            transform: translateY(-10px);
        }
    }
    
    .solvi-card, .solvi-metric, .solvi-inspiration-item, .diff-viewer {
        animation: fadeInUp 0.8s ease-out;
    }
    
    .solvi-logo {
        animation: pulse 5s ease-in-out infinite;
    }
    
    .solvi-badge {
        animation: float 6s ease-in-out infinite;
    }
    
    /* Scrollbar  personalizada */
    ::-webkit-scrollbar {
        width: 14px;
    }
    
    ::-webkit-scrollbar-track {
        background: var(--solvi-surface);
        border-radius: 14px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--solvi-gradient-primary);
        border-radius: 14px;
        border: 3px solid var(--solvi-surface);
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, var(--solvi-dark-green) 0%, var(--solvi-primary-green) 100%);
    }
    
    /* Efeitos especiais  */
    .solvi-glow {
        box-shadow: 0 0 20px var(--solvi-accent-green);
    }
    
    .solvi-shimmer {
        background: linear-gradient(45deg, transparent 30%, rgba(255,255,255,0.5) 50%, transparent 70%);
        background-size: 200% 200%;
        animation: shimmer 3s ease-in-out infinite;
    }
    
    @keyframes shimmer {
        0% {
            background-position: -200% -200%;
        }
        100% {
            background-position: 200% 200%;
        }
    }
</style>
""", unsafe_allow_html=True)

# Inicializar session state
def init_session_state():
    """Inicializa o estado da sessão com configurações """
    if 'current_tab' not in st.session_state:
        st.session_state.current_tab = 'comparison'  # Inicia com comparação para mostrar a nova funcionalidade
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = None
    if 'comparison_results' not in st.session_state:
        st.session_state.comparison_results = None
    if 'visual_diff_data' not in st.session_state:
        st.session_state.visual_diff_data = None

class FREAnalyzer:
    """Classe  para análise de FRE vs Normas CVM"""
    
    def __init__(self, api_key):
        openai.api_key = api_key
        self.client = openai.OpenAI(api_key=api_key)
        
    def extract_text_from_pdf(self, pdf_file):
        """Extrai texto de arquivo PDF com tratamento """
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extrai texto de arquivo Word com tratamento """
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do Word: {str(e)}")
            return ""
    
    def extract_text_from_file(self, uploaded_file):
        """Extrai texto baseado no tipo de arquivo"""
        if uploaded_file.type == "application/pdf":
            return self.extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                   "application/msword"]:
            return self.extract_text_from_docx(uploaded_file)
        else:
            st.error("❌ Formato de arquivo não suportado. Use PDF ou Word.")
            return ""
    
    def analyze_fre_section(self, fre_text, cvm_references, section_name, section_content):
        """Analisa uma seção específica do FRE contra as normas CVM"""
        
        prompt = f"""
        Você é um especialista  em regulamentação CVM e análise de Formulários de Referência (FRE).
        
        Analise a seção "{section_name}" do FRE fornecido contra as normas e orientações CVM.
        
        SEÇÃO ANALISADA:
        {section_content[:3000]}...
        
        NORMAS CVM DE REFERÊNCIA:
        {cvm_references[:5000]}...
        
        Para esta seção, identifique:
        
        1. CONFORMIDADE: Está em conformidade com as normas CVM?
        2. COMPLETUDE: Todas as informações obrigatórias estão presentes?
        3. QUALIDADE: A informação está clara, objetiva e completa?
        4. PONTOS DE ATENÇÃO: Identifique problemas específicos
        5. SUGESTÕES: Recomendações de melhoria com citação obrigatória dos artigos CVM
        
        RESPONDA EM JSON com esta estrutura:
        {{
            "secao": "{section_name}",
            "conformidade": "CONFORME/NAO_CONFORME/PARCIALMENTE_CONFORME",
            "criticidade": "CRITICO/ATENCAO/SUGESTAO",
            "pontos_atencao": [
                {{
                    "problema": "descrição do problema",
                    "criticidade": "CRITICO/ATENCAO/SUGESTAO",
                    "artigo_cvm": "artigo específico da norma CVM",
                    "sugestao": "recomendação específica de melhoria"
                }}
            ],
            "resumo": "resumo geral da análise desta seção"
        }}
        
        IMPORTANTE: 
        - Cite OBRIGATORIAMENTE os artigos específicos das normas CVM
        - Use criticidade CRITICO para não conformidades graves
        - Use ATENCAO para informações incompletas
        - Use SUGESTAO para melhorias recomendadas
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=2000
            )
            
            result = response.choices[0].message.content
            
            # Tenta extrair JSON da resposta
            try:
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                json_str = result[json_start:json_end]
                return json.loads(json_str)
            except:
                # Se falhar, retorna estrutura padrão
                return {
                    "secao": section_name,
                    "conformidade": "ERRO_ANALISE",
                    "criticidade": "ATENCAO",
                    "pontos_atencao": [{
                        "problema": "Erro na análise automática",
                        "criticidade": "ATENCAO",
                        "artigo_cvm": "Resolução CVM nº 80/22",
                        "sugestao": "Revisar manualmente esta seção"
                    }],
                    "resumo": "Erro na análise automática desta seção"
                }
                
        except Exception as e:
            st.error(f"❌ Erro na análise da seção {section_name}: {str(e)}")
            return None
    
    def extract_fre_sections(self, fre_text):
        """Extrai as seções principais do FRE"""
        sections = {}
        
        # Padrões para identificar seções do FRE
        section_patterns = [
            r"1\.1\s+Histórico do emissor",
            r"1\.2\s+Descrição das principais atividades",
            r"1\.3\s+Informações relacionadas aos segmentos operacionais",
            r"1\.4\s+Produção/Comercialização/Mercados",
            r"1\.5\s+Principais clientes",
            r"1\.6\s+Efeitos relevantes da regulação estatal",
            r"1\.9\s+Informações ambientais sociais e de governança",
            r"2\.1\s+Condições financeiras e patrimoniais",
            r"2\.2\s+Resultados operacional e financeiro",
            r"4\.1\s+Descrição dos fatores de risco",
            r"7\.1\s+Principais características dos órgãos de administração",
            r"8\.1\s+Política ou prática de remuneração",
            r"11\.1\s+Regras, políticas e práticas",
            r"12\.1\s+Informações sobre o capital social"
        ]
        
        # Divide o texto em seções
        lines = fre_text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            # Verifica se a linha corresponde a uma nova seção
            section_found = False
            for pattern in section_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    # Salva a seção anterior se existir
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Inicia nova seção
                    current_section = line.strip()
                    current_content = [line]
                    section_found = True
                    break
            
            if not section_found and current_section:
                current_content.append(line)
        
        # Salva a última seção
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections

class AdvancedDocumentComparator:
    """Classe  para comparação avançada de documentos com visualização"""
    
    def __init__(self):
        self.texto_ref = []
        self.texto_novo = []
        self.diferencas = []
        self.visual_diff_data = []
        
    def detectar_tipo_arquivo(self, nome_arquivo: str) -> str:
        """Detecta o tipo do arquivo baseado na extensão"""
        extensao = Path(nome_arquivo).suffix.lower()
        if extensao == '.pdf':
            return 'pdf'
        elif extensao in ['.docx', '.doc']:
            return 'word'
        else:
            return 'desconhecido'
    
    def extrair_texto_pdf(self, pdf_bytes: bytes) -> List[str]:
        """Extrai texto de cada página do PDF"""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            textos = []
            
            for i, pagina in enumerate(doc):
                texto = pagina.get_text()
                textos.append(texto)
            
            doc.close()
            return textos
            
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do PDF: {str(e)}")
            return []
    
    def extrair_texto_word(self, word_bytes: bytes) -> List[str]:
        """Extrai texto do documento Word"""
        try:
            # Salvar temporariamente para processar
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(word_bytes)
                tmp_path = tmp_file.name
            
            try:
                doc = docx.Document(tmp_path)
                
                # Para Word, vamos simular "páginas" agrupando parágrafos
                textos = []
                texto_atual = ""
                contador_paragrafos = 0
                
                for paragrafo in doc.paragraphs:
                    texto_atual += paragrafo.text + "\n"
                    contador_paragrafos += 1
                    
                    # Criar nova "página" a cada 50 parágrafos
                    if contador_paragrafos >= 50:
                        if texto_atual.strip():
                            textos.append(texto_atual)
                            texto_atual = ""
                            contador_paragrafos = 0
                
                # Adicionar último texto se houver
                if texto_atual.strip():
                    textos.append(texto_atual)
                
                # Se não há texto, criar pelo menos uma "página" vazia
                if not textos:
                    textos = [""]
                
                return textos
                
            finally:
                # Limpar arquivo temporário
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do Word: {str(e)}")
            return []
    
    def normalizar_texto_avancado(self, texto: str) -> str:
        """Normalização avançada de texto para comparação mais precisa"""
        # Remover espaços extras e quebras de linha desnecessárias
        texto = re.sub(r'\s+', ' ', texto.strip())
        
        # Remover caracteres de controle e formatação
        texto = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', texto)
        
        # Normalizar pontuação
        texto = re.sub(r'\s+([,.;:!?])', r'\1', texto)
        texto = re.sub(r'([,.;:!?])\s*', r'\1 ', texto)
        
        # Normalizar aspas e caracteres especiais
        texto = re.sub(r'["""]', '"', texto)
        texto = re.sub(r"[''']", "'", texto)
        texto = re.sub(r'[–—]', '-', texto)
        texto = re.sub(r'[…]', '...', texto)
        
        # Normalizar números e datas
        texto = re.sub(r'(\d+)\.(\d+)', r'\1,\2', texto)  # 1.000 -> 1,000
        texto = re.sub(r'(\d+)/(\d+)/(\d+)', r'\1-\2-\3', texto)  # 01/01/2024 -> 01-01-2024
        
        # Normalizar abreviações comuns
        abreviacoes = {
            r'\bSr\.\s*': 'Senhor ',
            r'\bSra\.\s*': 'Senhora ',
            r'\bDr\.\s*': 'Doutor ',
            r'\bDra\.\s*': 'Doutora ',
            r'\betc\.\s*': 'etcetera ',
            r'\bex\.\s*': 'exemplo ',
            r'\bp\.\s*ex\.\s*': 'por exemplo ',
        }
        
        for padrao, substituicao in abreviacoes.items():
            texto = re.sub(padrao, substituicao, texto, flags=re.IGNORECASE)
        
        return texto.strip()
    
    def dividir_em_sentencas_inteligente(self, texto: str) -> List[str]:
        """Divide o texto em sentenças de forma mais inteligente"""
        # Normalizar o texto primeiro
        texto = self.normalizar_texto_avancado(texto)
        
        # Substituir abreviações comuns primeiro para evitar quebras incorretas
        abreviacoes_protegidas = {
            'Sr.': 'Sr_TEMP_',
            'Sra.': 'Sra_TEMP_',
            'Dr.': 'Dr_TEMP_',
            'Dra.': 'Dra_TEMP_',
            'etc.': 'etc_TEMP_',
            'ex.': 'ex_TEMP_',
            'p.ex.': 'pex_TEMP_'
        }
        
        # Proteger abreviações
        for abrev, temp in abreviacoes_protegidas.items():
            texto = texto.replace(abrev, temp)
        
        # Dividir em sentenças usando padrão mais simples
        # Quebra em ponto seguido de espaço e letra maiúscula, mas não em números
        sentencas = re.split(r'\.(?!\d)\s+(?=[A-Z])', texto)
        
        # Restaurar abreviações
        sentencas_restauradas = []
        for sentenca in sentencas:
            for abrev, temp in abreviacoes_protegidas.items():
                sentenca = sentenca.replace(temp, abrev)
            sentencas_restauradas.append(sentenca)
        
        sentencas_limpas = []
        for sentenca in sentencas_restauradas:
            sentenca = sentenca.strip()
            if sentenca and len(sentenca) > 15:  # Filtrar sentenças muito curtas
                sentencas_limpas.append(sentenca)
        
        return sentencas_limpas
    
    def calcular_similaridade_avancada(self, texto1: str, texto2: str) -> float:
        """Calcula similaridade usando múltiplos algoritmos"""
        if not texto1 and not texto2:
            return 1.0
        if not texto1 or not texto2:
            return 0.0
        
        texto1_norm = self.normalizar_texto_avancado(texto1)
        texto2_norm = self.normalizar_texto_avancado(texto2)
        
        # Similaridade por sequência
        matcher = difflib.SequenceMatcher(None, texto1_norm, texto2_norm)
        sim_sequencia = matcher.ratio()
        
        # Similaridade por palavras
        palavras1 = set(texto1_norm.lower().split())
        palavras2 = set(texto2_norm.lower().split())
        
        if palavras1 or palavras2:
            intersecao = len(palavras1.intersection(palavras2))
            uniao = len(palavras1.union(palavras2))
            sim_palavras = intersecao / uniao if uniao > 0 else 0
        else:
            sim_palavras = 1.0
        
        # Média ponderada das similaridades
        similaridade_final = (sim_sequencia * 0.7) + (sim_palavras * 0.3)
        
        return similaridade_final
    
    def gerar_diff_visual_linha_por_linha(self, texto_ref: str, texto_novo: str) -> List[Dict]:
        """Gera diferenças visuais linha por linha para exibição"""
        linhas_ref = texto_ref.split('\n')
        linhas_novo = texto_novo.split('\n')
        
        # Usar difflib para comparação linha por linha
        differ = difflib.unified_diff(
            linhas_ref, 
            linhas_novo, 
            fromfile='Documento Original', 
            tofile='Documento Novo',
            lineterm=''
        )
        
        diff_lines = []
        linha_num = 1
        
        for line in differ:
            if line.startswith('@@'):
                # Cabeçalho de seção - ignorar
                continue
            elif line.startswith('---') or line.startswith('+++'):
                # Cabeçalho de arquivo - ignorar
                continue
            elif line.startswith('-'):
                # Linha removida
                diff_lines.append({
                    'numero': linha_num,
                    'tipo': 'removed',
                    'conteudo': line[1:],  # Remove o prefixo '-'
                    'conteudo_original': line[1:],
                    'conteudo_novo': ''
                })
                linha_num += 1
            elif line.startswith('+'):
                # Linha adicionada
                diff_lines.append({
                    'numero': linha_num,
                    'tipo': 'added',
                    'conteudo': line[1:],  # Remove o prefixo '+'
                    'conteudo_original': '',
                    'conteudo_novo': line[1:]
                })
                linha_num += 1
            elif line.startswith(' '):
                # Linha inalterada
                diff_lines.append({
                    'numero': linha_num,
                    'tipo': 'unchanged',
                    'conteudo': line[1:],  # Remove o prefixo ' '
                    'conteudo_original': line[1:],
                    'conteudo_novo': line[1:]
                })
                linha_num += 1
        
        return diff_lines
    
    def encontrar_alteracoes_avancadas(self, sentencas_ref: List[str], sentencas_novo: List[str]) -> List[Dict]:
        """Encontra alterações usando algoritmo avançado"""
        alteracoes = []
        
        # Criar conjuntos de sentenças únicas
        set_ref = set(sentencas_ref)
        set_novo = set(sentencas_novo)
        
        # Encontrar sentenças removidas e adicionadas
        sentencas_removidas = set_ref - set_novo
        sentencas_adicionadas = set_novo - set_ref
        
        # Verificar modificações usando similaridade avançada
        sentencas_modificadas = []
        
        for s_ref in sentencas_removidas.copy():
            melhor_match = None
            melhor_similaridade = 0.0
            
            for s_novo in sentencas_adicionadas:
                similaridade = self.calcular_similaridade_avancada(s_ref, s_novo)
                
                # Threshold mais baixo para detectar mais modificações
                if similaridade > 0.4 and similaridade > melhor_similaridade:
                    melhor_match = s_novo
                    melhor_similaridade = similaridade
            
            if melhor_match and melhor_similaridade > 0.4:
                sentencas_modificadas.append({
                    'original': s_ref,
                    'novo': melhor_match,
                    'similaridade': melhor_similaridade
                })
                sentencas_removidas.discard(s_ref)
                sentencas_adicionadas.discard(melhor_match)
        
        # Adicionar alterações
        for sentenca in sentencas_removidas:
            alteracoes.append({
                'tipo': 'removido',
                'texto': sentenca,
                'texto_original': sentenca,
                'texto_novo': '',
                'similaridade': 0.0
            })
        
        for sentenca in sentencas_adicionadas:
            alteracoes.append({
                'tipo': 'adicionado',
                'texto': sentenca,
                'texto_original': '',
                'texto_novo': sentenca,
                'similaridade': 0.0
            })
        
        for mod in sentencas_modificadas:
            alteracoes.append({
                'tipo': 'modificado',
                'texto': f"ANTES: {mod['original']}\nDEPOIS: {mod['novo']}",
                'texto_original': mod['original'],
                'texto_novo': mod['novo'],
                'similaridade': mod['similaridade']
            })
        
        return alteracoes

def render_header():
    """Renderiza o header masterpiece da aplicação"""
    st.markdown("""
    <div class="solvi-header">
        <div class="solvi-header-content">
            <div class="solvi-logo-section">
                <img src="https://static.wixstatic.com/media/b5b170_1e07cf7f7f82492a9808f9ae7f038596~mv2.png/v1/crop/x_0,y_0,w_2742,h_1106/fill/w_92,h_37,al_c,q_85,usm_0.66_1.00_0.01,enc_auto/Logotipo%20Solv%C3%AD_edited_edited.png" alt="Solví Logo" class="solvi-logo">
                <div class="solvi-title-section">
                    <h1 class="solvi-title">Plataforma Solví</h1>
                    <p class="solvi-subtitle">Análise Inteligente de Documentos com IA</p>
                </div>
            </div>
            <div class="solvi-badge">
                Soluções para a vida
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def render_inspiration_section():
    """Renderiza seção de imagens inspiracionais """
    st.markdown("""
    <div class="solvi-inspiration">
        <div class="solvi-inspiration-item">
            <img src="https://images.unsplash.com/photo-1558618666-fcd25c85cd64?ixlib=rb-4.0.3&auto=format&fit=crop&w=600&q=80" alt="Tecnologia Sustentável" class="solvi-inspiration-image">
            <div class="solvi-inspiration-content">
                <h3 class="solvi-inspiration-title">Tecnologia Sustentável</h3>
                <p class="solvi-inspiration-desc">Inovação em energia renovável e soluções tecnológicas verdes para um futuro sustentável e próspero para todas as gerações.</p>
            </div>
        </div>
        <div class="solvi-inspiration-item">
            <img src="https://images.unsplash.com/photo-1441974231531-c6227db76b6e?ixlib=rb-4.0.3&auto=format&fit=crop&w=600&q=80" alt="Proteção Ambiental" class="solvi-inspiration-image">
            <div class="solvi-inspiration-content">
                <h3 class="solvi-inspiration-title">Proteção Ambiental</h3>
                <p class="solvi-inspiration-desc">Preservação da natureza e biodiversidade através de práticas ambientais responsáveis e sustentáveis que protegem nosso planeta.</p>
            </div>
        </div>
        <div class="solvi-inspiration-item">
            <img src="https://images.unsplash.com/photo-1542601906990-b4d3fb778b09?ixlib=rb-4.0.3&auto=format&fit=crop&w=600&q=80" alt="Gestão de Resíduos" class="solvi-inspiration-image">
            <div class="solvi-inspiration-content">
                <h3 class="solvi-inspiration-title">Gestão de Resíduos</h3>
                <p class="solvi-inspiration-desc">Soluções inteligentes para reciclagem e economia circular, transformando resíduos em recursos valiosos para a sociedade.</p>
            </div>
        </div>
        <div class="solvi-inspiration-item">
            <img src="https://static.wixstatic.com/media/b5b170_b587909825174509a2a6a71af0106cc3~mv2.png/v1/fill/w_245,h_357,al_c,q_85,enc_auto/Group%2041.png" alt="Inovação Verde" class="solvi-inspiration-image">
            <div class="solvi-inspiration-content">
                <h3 class="solvi-inspiration-title">Inovação Verde</h3>
                <p class="solvi-inspiration-desc">Desenvolvimento de tecnologias limpas e processos inovadores para sustentabilidade empresarial e crescimento responsável.</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def render_navigation():
    """Renderiza a navegação  por abas"""
    st.markdown('<div class="solvi-navigation">', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Análise CVM", key="tab_cvm", use_container_width=True):
            st.session_state.current_tab = 'cvm'
            st.rerun()
    
    with col2:
        if st.button("Comparação Visual de Documentos", key="tab_comparison", use_container_width=True):
            st.session_state.current_tab = 'comparison'
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)

def render_visual_diff_viewer(diff_data: List[Dict], arquivo_ref: str, arquivo_novo: str):
    """Renderiza o visualizador avançado de diferenças usando componentes Streamlit nativos"""
    
    # Calcular estatísticas
    total_lines = len(diff_data)
    added_lines = len([d for d in diff_data if d['tipo'] == 'added'])
    removed_lines = len([d for d in diff_data if d['tipo'] == 'removed'])
    modified_lines = len([d for d in diff_data if d['tipo'] == 'modified'])
    unchanged_lines = total_lines - added_lines - removed_lines - modified_lines
    
    # Header com estatísticas usando Streamlit nativo
    st.markdown(f"""
    <div class="diff-viewer">
        <div class="diff-header">
            <h3 class="diff-title">Comparação Visual: {arquivo_ref} ↔ {arquivo_novo}</h3>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Estatísticas em colunas
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            label="Adições",
            value=added_lines,
            delta=f"{(added_lines/total_lines*100):.1f}%" if total_lines > 0 else "0%"
        )
    
    with col2:
        st.metric(
            label="Remoções", 
            value=removed_lines,
            delta=f"{(removed_lines/total_lines*100):.1f}%" if total_lines > 0 else "0%"
        )
    
    with col3:
        st.metric(
            label="Modificações",
            value=modified_lines,
            delta=f"{(modified_lines/total_lines*100):.1f}%" if total_lines > 0 else "0%"
        )
    
    with col4:
        st.metric(
            label="Total de Linhas",
            value=total_lines,
            delta=f"Taxa de mudança: {((added_lines + removed_lines + modified_lines) / total_lines * 100):.1f}%" if total_lines > 0 else "0%"
        )
    
    # Controles de visualização aprimorados
    st.markdown("### Controles de Visualização")
    
    col_ctrl1, col_ctrl2, col_ctrl3 = st.columns(3)
    
    with col_ctrl1:
        filtro_tipo = st.selectbox(
            "Filtrar por tipo:",
            ["Todos", "Apenas adicionados", "Apenas removidos", "Apenas modificados", "Apenas inalterados", "Apenas diferenças"],
            index=0
        )
    
    with col_ctrl2:
        max_linhas = st.selectbox(
            "Máximo de linhas:",
            [50, 100, 200, 500, 1000],
            index=2
        )
    
    with col_ctrl3:
        view_mode = st.radio(
            "Modo de visualização:",
            ["Unificado", "Lado a lado"],
            index=0,
            horizontal=True
        )
    
    # Filtrar dados baseado na seleção
    if filtro_tipo == "Apenas adicionados":
        filtered_data = [d for d in diff_data if d['tipo'] == 'added']
    elif filtro_tipo == "Apenas removidos":
        filtered_data = [d for d in diff_data if d['tipo'] == 'removed']
    elif filtro_tipo == "Apenas modificados":
        filtered_data = [d for d in diff_data if d['tipo'] == 'modified']
    elif filtro_tipo == "Apenas inalterados":
        filtered_data = [d for d in diff_data if d['tipo'] == 'unchanged']
    elif filtro_tipo == "Apenas diferenças":
        filtered_data = [d for d in diff_data if d['tipo'] != 'unchanged']
    else:  # Todos
        filtered_data = diff_data
    
    # Limitar linhas baseado na seleção
    display_data = filtered_data[:max_linhas]
    
    if len(filtered_data) > max_linhas:
        st.warning(f"Mostrando apenas as primeiras {max_linhas} linhas de {len(filtered_data)} total. Ajuste o filtro para ver mais.")
    
    # Informações de filtro aplicado
    st.info(f"Filtro aplicado: {filtro_tipo} | Exibindo: {len(display_data)} de {len(filtered_data)} linhas filtradas | Total no documento: {total_lines} linhas")
    
    # Visualização das diferenças
    st.markdown("### Diferenças Detectadas")
    
    # Container com scroll para as diferenças
    diff_container = st.container()
    
    with diff_container:
        for i, line_data in enumerate(display_data):
            tipo = line_data['tipo']
            numero = line_data['numero']
            conteudo = line_data['conteudo'][:500]  # Aumentar limite para mais contexto
            
            # Calcular página aproximada (assumindo 50 linhas por página)
            pagina = (numero // 50) + 1
            linha_na_pagina = numero % 50 if numero % 50 != 0 else 50
            
            # Escolher cor e ícone baseado no tipo
            if tipo == 'added':
                st.success(f"**Linha {numero}** (Página {pagina}, Linha {linha_na_pagina}) | **ADICIONADA**\n```\n{conteudo}\n```")
            elif tipo == 'removed':
                st.error(f"**Linha {numero}** (Página {pagina}, Linha {linha_na_pagina}) | **REMOVIDA**\n```\n{conteudo}\n```")
            elif tipo == 'modified':
                st.warning(f"**Linha {numero}** (Página {pagina}, Linha {linha_na_pagina}) | **MODIFICADA**\n```\n{conteudo}\n```")
            else:  # unchanged
                st.info(f"**Linha {numero}** (Página {pagina}, Linha {linha_na_pagina}) | **INALTERADA**\n```\n{conteudo}\n```")
    
    # Resumo final
    st.markdown("### Resumo da Análise")
    
    resumo_cols = st.columns(6)
    
    with resumo_cols[0]:
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{total_lines}</div>
            <div class="solvi-metric-label">Total de Linhas</div>
        </div>
        """, unsafe_allow_html=True)
    
    with resumo_cols[1]:
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{unchanged_lines}</div>
            <div class="solvi-metric-label">Inalteradas</div>
        </div>
        """, unsafe_allow_html=True)
    
    with resumo_cols[2]:
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{added_lines}</div>
            <div class="solvi-metric-label">Adicionadas</div>
        </div>
        """, unsafe_allow_html=True)
    
    with resumo_cols[3]:
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{removed_lines}</div>
            <div class="solvi-metric-label">Removidas</div>
        </div>
        """, unsafe_allow_html=True)
    
    with resumo_cols[4]:
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{modified_lines}</div>
            <div class="solvi-metric-label">Modificadas</div>
        </div>
        """, unsafe_allow_html=True)
    
    with resumo_cols[5]:
        taxa_mudanca = ((added_lines + removed_lines + modified_lines) / total_lines * 100) if total_lines > 0 else 0
        st.markdown(f"""
        <div class="solvi-metric">
            <div class="solvi-metric-value">{taxa_mudanca:.1f}%</div>
            <div class="solvi-metric-label">Taxa de Mudança</div>
        </div>
        """, unsafe_allow_html=True)

def render_cvm_analysis():
    """Renderiza a interface  de análise CVM com SIDEBAR CORRIGIDA"""
    st.markdown("""
    <div class="solvi-card">
        <div class="solvi-card-header">
            <div class="solvi-card-icon">📊</div>
            <h2 class="solvi-card-title">Análise FRE vs Normas CVM</h2>
        </div>
        <p style="color: #666; font-size: 1.4rem; line-height: 1.8; font-weight: 500; font-family: 'Inter', sans-serif;">
            Análise automatizada  de Formulários de Referência contra normas CVM com identificação 
            inteligente de não conformidades e geração de relatórios detalhados com base legal específica 
            e recomendações de melhoria personalizadas.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # SIDEBAR  CORRIGIDA - SEMPRE VISÍVEL
    with st.sidebar:
        st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
        st.markdown('<h3 class="sidebar-title">⚙️ Configurações </h3>', unsafe_allow_html=True)
        
        # Campo obrigatório para API Key
        api_key = st.text_input(
            "🔑 Chave API OpenAI *",
            type="password",
            help="Insira sua chave API da OpenAI (obrigatório para análise )"
        )
        
        if not api_key:
            st.markdown("""
            <div class="solvi-alert error">
                ⚠️ <strong>Chave API OpenAI é obrigatória!</strong><br>
                Configure sua chave para utilizar a análise CVM  com IA avançada.
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Upload do FRE
        st.markdown("### 📄 Arquivo FRE")
        fre_file = st.file_uploader(
            "Upload do Formulário de Referência",
            type=['pdf', 'docx'],
            help="Faça upload do FRE para análise "
        )
        
        st.markdown("---")
        
        # Upload dos documentos CVM
        st.markdown("### 📚 Documentos CVM (máx. 5)")
        cvm_files = st.file_uploader(
            "Upload dos documentos de referência CVM",
            type=['pdf', 'docx'],
            accept_multiple_files=True,
            help="Faça upload dos documentos CVM para comparação avançada"
        )
        
        if len(cvm_files) > 5:
            st.error("⚠️ Máximo de 5 documentos CVM permitidos para análise !")
            cvm_files = cvm_files[:5]
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Área principal 
    if not api_key:
        st.markdown("""
        <div class="solvi-upload">
            <div class="solvi-upload-icon">🔑</div>
            <div class="solvi-upload-text">Configure sua API Key OpenAI</div>
            <div class="solvi-upload-subtext">
                Para utilizar a análise CVM , você precisa configurar sua chave API OpenAI na barra lateral.<br>
                A chave é necessária para processar documentos com inteligência artificial avançada.
            </div>
        </div>
        """, unsafe_allow_html=True)
        return
    
    if not fre_file:
        st.markdown("""
        <div class="solvi-upload">
            <div class="solvi-upload-icon">📄</div>
            <div class="solvi-upload-text">Como usar a Análise CVM </div>
            <div class="solvi-upload-subtext">
                1. Configure sua API Key OpenAI na barra lateral<br>
                2. Faça upload do FRE (Formulário de Referência)<br>
                3. Adicione documentos CVM para comparação avançada<br>
                4. Execute a análise  e receba relatório detalhado<br>
                5. Baixe relatórios em PDF com recomendações personalizadas
            </div>
        </div>
        """, unsafe_allow_html=True)
        return
    
    if not cvm_files:
        st.markdown("""
        <div class="solvi-alert warning">
            ⚠️ <strong>Documentos CVM necessários para análise </strong><br>
            Adicione pelo menos um documento CVM para realizar a análise comparativa avançada com IA.
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Informações  dos arquivos carregados
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"""
        <div class="solvi-alert success">
            ✅ <strong>FRE Carregado:</strong> {fre_file.name}<br>
            📊 <strong>Tamanho:</strong> {fre_file.size / 1024 / 1024:.2f} MB<br>
            🎯 <strong>Status:</strong> Pronto para análise 
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="solvi-alert info">
            📚 <strong>Documentos CVM:</strong> {len(cvm_files)} arquivo(s)<br>
            📊 <strong>Total:</strong> {sum(f.size for f in cvm_files) / 1024 / 1024:.2f} MB<br>
            🚀 <strong>Análise:</strong> IA avançada habilitada
        </div>
        """, unsafe_allow_html=True)
    
    # Botão de análise 
    if st.button("🔍 Iniciar Análise CVM ", type="primary", use_container_width=True):
        with st.spinner("🔄 Processando análise  com IA..."):
            try:
                # Inicializar analisador 
                analyzer = FREAnalyzer(api_key)
                
                # Extrair texto do FRE
                fre_text = analyzer.extract_text_from_file(fre_file)
                if not fre_text:
                    st.error("❌ Erro ao extrair texto do FRE")
                    return
                
                # Extrair texto dos documentos CVM
                cvm_text = ""
                for cvm_file in cvm_files:
                    cvm_content = analyzer.extract_text_from_file(cvm_file)
                    cvm_text += cvm_content + "\n\n"
                
                if not cvm_text:
                    st.error("❌ Erro ao extrair texto dos documentos CVM")
                    return
                
                # Extrair seções do FRE
                fre_sections = analyzer.extract_fre_sections(fre_text)
                
                if not fre_sections:
                    st.error("❌ Não foi possível identificar seções no FRE")
                    return
                
                # Analisar cada seção com IA 
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                analysis_results = []
                total_sections = len(fre_sections)
                
                for i, (section_name, section_content) in enumerate(fre_sections.items()):
                    status_text.text(f"🤖 Analisando com IA: {section_name}")
                    
                    result = analyzer.analyze_fre_section(
                        fre_text, cvm_text, section_name, section_content
                    )
                    
                    if result:
                        analysis_results.append(result)
                    
                    progress_bar.progress((i + 1) / total_sections)
                    time.sleep(0.5)
                
                status_text.text("✅ Análise  concluída com sucesso!")
                progress_bar.empty()
                status_text.empty()
                
                # Salvar resultados 
                st.session_state.analysis_results = analysis_results
                st.session_state.fre_filename = fre_file.name
                
                st.markdown("""
                <div class="solvi-alert success">
                    ✅ <strong>Análise CVM  concluída com sucesso!</strong><br>
                    Confira os resultados detalhados e insights avançados abaixo.
                </div>
                """, unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"❌ Erro durante a análise : {str(e)}")
    
    # Exibir resultados  se disponíveis
    if st.session_state.analysis_results:
        analysis_results = st.session_state.analysis_results
        
        st.markdown("### 📊 Resultados da Análise ")
        
        # Métricas 
        total_pontos = sum(len(r.get('pontos_atencao', [])) for r in analysis_results)
        criticos = sum(1 for r in analysis_results for p in r.get('pontos_atencao', []) if p.get('criticidade') == 'CRITICO')
        atencao = sum(1 for r in analysis_results for p in r.get('pontos_atencao', []) if p.get('criticidade') == 'ATENCAO')
        sugestoes = sum(1 for r in analysis_results for p in r.get('pontos_atencao', []) if p.get('criticidade') == 'SUGESTAO')
        
        st.markdown(f"""
        <div class="solvi-metrics">
            <div class="solvi-metric">
                <div class="solvi-metric-value">{total_pontos}</div>
                <div class="solvi-metric-label">Total de Pontos</div>
            </div>
            <div class="solvi-metric">
                <div class="solvi-metric-value">{criticos}</div>
                <div class="solvi-metric-label">Críticos</div>
            </div>
            <div class="solvi-metric">
                <div class="solvi-metric-value">{atencao}</div>
                <div class="solvi-metric-label">Atenção</div>
            </div>
            <div class="solvi-metric">
                <div class="solvi-metric-value">{sugestoes}</div>
                <div class="solvi-metric-label">Sugestões</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Exibir resultados detalhados 
        for result in analysis_results:
            with st.expander(f"📑 {result.get('secao', 'Seção não identificada')}", expanded=False):
                conformidade = result.get('conformidade', 'N/A')
                if conformidade == 'CONFORME':
                    st.success(f"✅ Status: {conformidade}")
                elif conformidade == 'NAO_CONFORME':
                    st.error(f"❌ Status: {conformidade}")
                else:
                    st.warning(f"⚠️ Status: {conformidade}")
                
                st.write(f"**Resumo:** {result.get('resumo', 'N/A')}")
                
                pontos = result.get('pontos_atencao', [])
                if pontos:
                    st.write("**Pontos de Atenção:**")
                    for i, ponto in enumerate(pontos, 1):
                        criticidade = ponto.get('criticidade', 'N/A')
                        emoji = "🔴" if criticidade == "CRITICO" else "🟡" if criticidade == "ATENCAO" else "🟢"
                        
                        st.write(f"{emoji} **Ponto {i}:** {ponto.get('problema', 'N/A')}")
                        st.write(f"**Base legal:** {ponto.get('artigo_cvm', 'N/A')}")
                        st.write(f"**Sugestão:** {ponto.get('sugestao', 'N/A')}")
                        st.write("---")

def render_document_comparison():
    """Renderiza a interface  de comparação visual de documentos"""
    st.markdown("""
    <div class="solvi-card">
        <div class="solvi-card-header">
            <div class="solvi-card-icon">📚</div>
            <h2 class="solvi-card-title">Comparação Visual Avançada de Documentos</h2>
        </div>
        <p style="color: #666; font-size: 1.4rem; line-height: 1.8; font-weight: 500; font-family: 'Inter', sans-serif;">
            Compare dois documentos (PDF ou Word) com algoritmo  de IA e visualize as diferenças 
            de forma interativa, com destaque visual para adições, remoções e modificações linha por linha, 
            similar ao GitHub ou Google Docs.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Layout  em colunas para upload
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Documento de Referência")
        arquivo_ref = st.file_uploader(
            "Escolha o arquivo de referência",
            type=['pdf', 'docx'],
            key="ref_uploader",
            help="Este será usado como base para comparação "
        )
        
        if arquivo_ref:
            st.markdown(f"""
            <div class="solvi-alert success">
                ✅ <strong>Arquivo carregado:</strong> {arquivo_ref.name}<br>
                📊 <strong>Tamanho:</strong> {arquivo_ref.size / 1024 / 1024:.2f} MB<br>
                📋 <strong>Tipo:</strong> {arquivo_ref.type.split('/')[-1].upper()}<br>
                🎯 <strong>Status:</strong> Pronto para análise 
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("### 📄 Novo Documento")
        arquivo_novo = st.file_uploader(
            "Escolha o novo arquivo",
            type=['pdf', 'docx'],
            key="novo_uploader",
            help="Este será comparado com o arquivo de referência"
        )
        
        if arquivo_novo:
            st.markdown(f"""
            <div class="solvi-alert success">
                ✅ <strong>Arquivo carregado:</strong> {arquivo_novo.name}<br>
                📊 <strong>Tamanho:</strong> {arquivo_novo.size / 1024 / 1024:.2f} MB<br>
                📋 <strong>Tipo:</strong> {arquivo_novo.type.split('/')[-1].upper()}<br>
                🚀 <strong>Análise:</strong> Visualização avançada habilitada
            </div>
            """, unsafe_allow_html=True)
    
    # Informações  sobre o algoritmo com pontilhado contínuo
    if not arquivo_ref or not arquivo_novo:
        st.markdown("""
        <div class="solvi-upload">
            <div class="solvi-upload-icon">🎨</div>
            <div class="solvi-upload-text">Visualização Avançada de Diferenças</div>
            <div class="solvi-upload-subtext">
                ✅ Comparação visual linha por linha como GitHub<br>
                ✅ Destaque colorido para adições, remoções e modificações<br>
                ✅ Algoritmo avançado de normalização de texto<br>
                ✅ Estatísticas detalhadas de alterações<br>
                ✅ Interface interativa com controles de visualização<br>
                ✅ Relatórios  com insights visuais
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Botão de comparação 
    if arquivo_ref and arquivo_novo:
        # Verificar compatibilidade de tipos
        comparator = AdvancedDocumentComparator()
        tipo_ref = comparator.detectar_tipo_arquivo(arquivo_ref.name)
        tipo_novo = comparator.detectar_tipo_arquivo(arquivo_novo.name)
        
        if tipo_ref != tipo_novo:
            st.markdown(f"""
            <div class="solvi-alert warning">
                ⚠️ <strong>Tipos diferentes detectados:</strong> {tipo_ref.upper()} vs {tipo_novo.upper()}<br>
                A comparação visual ainda é possível com algoritmo adaptativo, mas pode não ser ideal.
            </div>
            """, unsafe_allow_html=True)
        
        if st.button("🎨 Comparar com Visualização Avançada", type="primary", use_container_width=True):
            with st.spinner("🔄 Processando comparação visual ..."):
                try:
                    # Extrair textos 
                    ref_bytes = arquivo_ref.read()
                    novo_bytes = arquivo_novo.read()
                    
                    if tipo_ref == 'pdf':
                        texto_ref_pages = comparator.extrair_texto_pdf(ref_bytes)
                    else:
                        texto_ref_pages = comparator.extrair_texto_word(ref_bytes)
                    
                    if tipo_novo == 'pdf':
                        texto_novo_pages = comparator.extrair_texto_pdf(novo_bytes)
                    else:
                        texto_novo_pages = comparator.extrair_texto_word(novo_bytes)
                    
                    if not texto_ref_pages or not texto_novo_pages:
                        st.error("❌ Erro ao extrair texto dos documentos")
                        return
                    
                    # Combinar todas as páginas em um texto único
                    texto_ref_completo = '\n'.join(texto_ref_pages)
                    texto_novo_completo = '\n'.join(texto_novo_pages)
                    
                    # Gerar diferenças visuais linha por linha
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("🎨 Gerando visualização linha por linha...")
                    progress_bar.progress(0.3)
                    
                    diff_visual = comparator.gerar_diff_visual_linha_por_linha(
                        texto_ref_completo, texto_novo_completo
                    )
                    
                    progress_bar.progress(0.6)
                    status_text.text("📊 Calculando estatísticas avançadas...")
                    
                    # Comparar textos com algoritmo avançado para estatísticas
                    sentencas_ref = comparator.dividir_em_sentencas_inteligente(texto_ref_completo)
                    sentencas_novo = comparator.dividir_em_sentencas_inteligente(texto_novo_completo)
                    
                    alteracoes_avancadas = comparator.encontrar_alteracoes_avancadas(
                        sentencas_ref, sentencas_novo
                    )
                    
                    progress_bar.progress(1.0)
                    status_text.text("✅ Visualização  concluída!")
                    
                    time.sleep(0.5)
                    progress_bar.empty()
                    status_text.empty()
                    
                    # Salvar resultados 
                    st.session_state.visual_diff_data = diff_visual
                    st.session_state.comparison_results = {
                        'diferencas': alteracoes_avancadas,
                        'arquivo_ref': arquivo_ref.name,
                        'arquivo_novo': arquivo_novo.name,
                        'diff_visual': diff_visual
                    }
                    
                    st.markdown("""
                    <div class="solvi-alert success">
                        ✅ <strong>Comparação visual  concluída com sucesso!</strong><br>
                        Confira a visualização avançada e insights detalhados abaixo.
                    </div>
                    """, unsafe_allow_html=True)
                    
                except Exception as e:
                    st.error(f"❌ Erro durante a comparação visual: {str(e)}")
    
    # Exibir visualização  se disponível
    if st.session_state.visual_diff_data:
        st.markdown("### 🎨 Visualização Avançada de Diferenças")
        
        # Renderizar o visualizador 
        render_visual_diff_viewer(
            st.session_state.visual_diff_data,
            st.session_state.comparison_results['arquivo_ref'],
            st.session_state.comparison_results['arquivo_novo']
        )
        
        # Tabela de alterações semânticas
        if st.session_state.comparison_results['diferencas']:
            st.markdown("### 📋 Análise Semântica Detalhada")
            
            alteracoes = st.session_state.comparison_results['diferencas']
            
            # Métricas semânticas
            total_alteracoes = len(alteracoes)
            adicionados = len([a for a in alteracoes if a['tipo'] == 'adicionado'])
            removidos = len([a for a in alteracoes if a['tipo'] == 'removido'])
            modificados = len([a for a in alteracoes if a['tipo'] == 'modificado'])
            
            st.markdown(f"""
            <div class="solvi-metrics">
                <div class="solvi-metric">
                    <div class="solvi-metric-value">{total_alteracoes}</div>
                    <div class="solvi-metric-label">Alterações Semânticas</div>
                </div>
                <div class="solvi-metric">
                    <div class="solvi-metric-value">{adicionados}</div>
                    <div class="solvi-metric-label">Sentenças Adicionadas</div>
                </div>
                <div class="solvi-metric">
                    <div class="solvi-metric-value">{removidos}</div>
                    <div class="solvi-metric-label">Sentenças Removidas</div>
                </div>
                <div class="solvi-metric">
                    <div class="solvi-metric-value">{modificados}</div>
                    <div class="solvi-metric-label">Sentenças Modificadas</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Tabela de alterações
            df_alteracoes = []
            for i, alteracao in enumerate(alteracoes, 1):
                df_alteracoes.append({
                    'ID': i,
                    'Tipo': alteracao['tipo'].title(),
                    'Texto Original': alteracao['texto_original'][:100] + '...' if len(alteracao['texto_original']) > 100 else alteracao['texto_original'],
                    'Texto Novo': alteracao['texto_novo'][:100] + '...' if len(alteracao['texto_novo']) > 100 else alteracao['texto_novo'],
                    'Similaridade': f"{alteracao.get('similaridade', 0):.2%}" if alteracao.get('similaridade') else 'N/A'
                })
            
            df = pd.DataFrame(df_alteracoes)
            st.dataframe(df, use_container_width=True)

def render_footer():
    """Renderiza o footer  da aplicação"""
    st.markdown("""
    <div class="solvi-footer">
        <div class="solvi-footer-content">
            <img src="https://static.wixstatic.com/media/b5b170_1e07cf7f7f82492a9808f9ae7f038596~mv2.png/v1/crop/x_0,y_0,w_2742,h_1106/fill/w_92,h_37,al_c,q_85,usm_0.66_1.00_0.01,enc_auto/Logotipo%20Solv%C3%AD_edited_edited.png" alt="Solví Logo" class="solvi-footer-logo">
            <p style="margin: 3rem 0 1.5rem 0; font-size: 1.6rem; font-weight: 800; font-family: 'Poppins', sans-serif;">
                🌱 Plataforma Solví - Soluções Inteligentes  para Análise de Documentos
            </p>
            <p style="margin: 0; opacity: 0.95; font-size: 1.2rem; font-weight: 600; font-family: 'Inter', sans-serif;">
                Desenvolvido com ❤️ para sustentabilidade e inovação • Soluções para a vida • Tecnologia 
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

def main():
    """Função principal masterpiece da aplicação"""
    # Inicializar session state 
    init_session_state()
    
    # Renderizar header masterpiece
    render_header()
    
    # Renderizar seção inspiracional 
    render_inspiration_section()
    
    # Renderizar navegação 
    render_navigation()
    
    # Renderizar conteúdo  baseado na aba selecionada
    if st.session_state.current_tab == 'cvm':
        render_cvm_analysis()
    elif st.session_state.current_tab == 'comparison':
        render_document_comparison()
    else:
        # Fallback para comparação para mostrar a nova funcionalidade
        st.session_state.current_tab = 'comparison'
        render_document_comparison()
    
    # Renderizar footer 
    render_footer()

if __name__ == "__main__":
    main()
