"""
Aplicação Streamlit Unificada - Solvi
Combina Comparador de Documentos e Analisador FRE vs CVM
"""

import streamlit as st
import fitz  # PyMuPDF
import difflib
import pandas as pd
import io
from datetime import datetime
import base64
from typing import List, Tuple, Dict, Optional, Set
import logging
from pathlib import Path
import tempfile
import os
import re
import openai
import PyPDF2
import docx
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import json
import time

# Configuração da página
st.set_page_config(
    page_title="Plataforma Solvi - Análise de Documentos",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Inicializar session state
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
if 'comparison_results' not in st.session_state:
    st.session_state.comparison_results = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'fre_filename' not in st.session_state:
    st.session_state.fre_filename = None
if 'analysis_completed' not in st.session_state:
    st.session_state.analysis_completed = False

# CSS customizado com cores do logo
st.markdown("""
<style>
    /* Reset e configurações gerais */
    .main {
        background-color: #ffffff;
        color: #000000;
    }
    
    .stApp {
        background-color: #ffffff;
    }
    
    /* Header com logo */
    .header-container {
        display: flex;
        align-items: center;
        justify-content: space-between;
        background: linear-gradient(135deg, #1e3a8a 0%, #2563eb 100%);
        padding: 20px 30px;
        border-radius: 15px;
        margin-bottom: 30px;
        box-shadow: 0 8px 32px rgba(30, 58, 138, 0.2);
    }
    
    .header-text {
        color: white;
        flex: 1;
    }
    
    .header-title {
        font-size: 2.5em;
        font-weight: bold;
        margin: 0;
        color: white;
    }
    
    .header-subtitle {
        font-size: 1.1em;
        margin: 5px 0 0 0;
        opacity: 0.9;
        color: white;
    }
    
    .logo-container {
        background: white;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    .logo-container img {
        max-height: 60px;
        max-width: 150px;
        object-fit: contain;
    }
    
    /* Abas customizadas */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: #f8fafc;
        padding: 8px;
        border-radius: 10px;
        border: 1px solid #e2e8f0;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: white;
        border-radius: 8px;
        color: #1e3a8a;
        font-weight: 600;
        border: 1px solid #e2e8f0;
        transition: all 0.3s ease;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #1e3a8a !important;
        color: white !important;
        border-color: #1e3a8a !important;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #f1f5f9;
        border-color: #2563eb;
    }
    
    .stTabs [aria-selected="true"]:hover {
        background-color: #1e40af !important;
    }
    
    /* Botões */
    .stButton > button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.5rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        background-color: #1e40af;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(30, 58, 138, 0.3);
    }
    
    /* Caixas de informação */
    .info-box {
        background-color: #eff6ff;
        border: 1px solid #2563eb;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        color: #1e40af;
    }
    
    .warning-box {
        background-color: #fef3c7;
        border: 1px solid #f59e0b;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        color: #92400e;
    }
    
    .success-box {
        background-color: #f0fdf4;
        border: 1px solid #22c55e;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        color: #166534;
    }
    
    /* Métricas */
    .metric-container {
        background: white;
        padding: 20px;
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .metric-value {
        font-size: 2em;
        font-weight: bold;
        color: #1e3a8a;
        margin-bottom: 5px;
    }
    
    .metric-label {
        color: #64748b;
        font-size: 0.9em;
    }
    
    /* Parágrafos para comparação */
    .paragrafo-container {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        margin: 15px 0;
        overflow: hidden;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .paragrafo-header {
        background: linear-gradient(135deg, #1e3a8a 0%, #2563eb 100%);
        color: white;
        padding: 12px 20px;
        font-weight: bold;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .paragrafo-content {
        padding: 20px;
        font-family: 'Georgia', 'Times New Roman', serif;
        font-size: 14px;
        line-height: 1.8;
        background: #fafafa;
        color: #000000;
    }
    
    .paragrafo-adicionado {
        background-color: #f0fdf4;
        border-left: 4px solid #22c55e;
        color: #166534;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-removido {
        background-color: #fef2f2;
        border-left: 4px solid #ef4444;
        color: #dc2626;
        text-decoration: line-through;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-modificado {
        background-color: #fffbeb;
        border-left: 4px solid #f59e0b;
        color: #92400e;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    .paragrafo-normal {
        background-color: #f8fafc;
        border-left: 4px solid #e2e8f0;
        color: #475569;
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 4px;
    }
    
    /* Sidebar */
    .css-1d391kg {
        background-color: #f8fafc;
    }
    
    /* Inputs */
    .stTextInput > div > div > input {
        background-color: white;
        color: #000000;
        border: 1px solid #e2e8f0;
    }
    
    .stSelectbox > div > div > select {
        background-color: white;
        color: #000000;
        border: 1px solid #e2e8f0;
    }
    
    /* File uploader */
    .stFileUploader > div {
        background-color: white;
        border: 2px dashed #2563eb;
        border-radius: 8px;
    }
    
    /* Progress bar */
    .stProgress > div > div > div > div {
        background-color: #1e3a8a;
    }
    
    /* Responsividade */
    @media (max-width: 768px) {
        .header-container {
            flex-direction: column;
            text-align: center;
            gap: 20px;
        }
        
        .header-title {
            font-size: 2em;
        }
        
        .logo-container {
            order: -1;
        }
    }
</style>
""", unsafe_allow_html=True)

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Função para carregar logo
@st.cache_data
def load_logo():
    try:
        with open("logo.png", "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

# Header com logo
def render_header():
    logo_base64 = load_logo()
    
    if logo_base64:
        st.markdown(f"""
        <div class="header-container">
            <div class="header-text">
                <h1 class="header-title">Plataforma Solvi</h1>
                <p class="header-subtitle">Análise Inteligente de Documentos</p>
            </div>
            <div class="logo-container">
                <img src="data:image/png;base64,{logo_base64}" alt="Logo Solvi">
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="header-container">
            <div class="header-text">
                <h1 class="header-title">Plataforma Solvi</h1>
                <p class="header-subtitle">Análise Inteligente de Documentos</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

# Classe para comparação de documentos
class DocumentComparator:
    """Classe principal para comparação de documentos (PDF e Word)"""
    
    def __init__(self):
        self.texto_ref = []
        self.texto_novo = []
        self.diferencas = []
        self.diferencas_detalhadas = []
        self.tipo_ref = None
        self.tipo_novo = None
        
    def detectar_tipo_arquivo(self, nome_arquivo: str) -> str:
        """Detecta o tipo do arquivo baseado na extensão"""
        extensao = Path(nome_arquivo).suffix.lower()
        if extensao == '.pdf':
            return 'pdf'
        elif extensao in ['.docx', '.doc']:
            return 'word'
        else:
            return 'desconhecido'
    
    def validar_arquivo(self, arquivo_bytes: bytes, nome_arquivo: str) -> bool:
        """Valida se o arquivo é válido baseado no tipo"""
        tipo = self.detectar_tipo_arquivo(nome_arquivo)
        
        try:
            if tipo == 'pdf':
                doc = fitz.open(stream=arquivo_bytes, filetype="pdf")
                if doc.page_count == 0:
                    st.error(f"❌ O arquivo PDF '{nome_arquivo}' não contém páginas.")
                    return False
                doc.close()
                return True
                
            elif tipo == 'word':
                try:
                    from docx import Document
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                        tmp_file.write(arquivo_bytes)
                        tmp_path = tmp_file.name
                    
                    try:
                        doc = Document(tmp_path)
                        if len(doc.paragraphs) == 0:
                            st.error(f"❌ O arquivo Word '{nome_arquivo}' não contém texto.")
                            return False
                        return True
                    except Exception as e:
                        st.error(f"❌ Erro ao abrir arquivo Word '{nome_arquivo}': {str(e)}")
                        return False
                    finally:
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                except ImportError:
                    st.error("❌ Biblioteca python-docx não está disponível. Instale com: pip install python-docx")
                    return False
            else:
                st.error(f"❌ Tipo de arquivo não suportado: {nome_arquivo}")
                return False
                
        except Exception as e:
            st.error(f"❌ Erro ao validar '{nome_arquivo}': {str(e)}")
            return False
    
    def extrair_texto_pdf(self, pdf_bytes: bytes) -> List[str]:
        """Extrai texto de cada página do PDF"""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            textos = []
            
            for i, pagina in enumerate(doc):
                texto = pagina.get_text()
                textos.append(texto)
            
            doc.close()
            return textos
            
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do PDF: {str(e)}")
            return []
    
    def extrair_texto_word(self, word_bytes: bytes) -> List[str]:
        """Extrai texto do documento Word"""
        try:
            from docx import Document
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(word_bytes)
                tmp_path = tmp_file.name
            
            try:
                doc = Document(tmp_path)
                textos = []
                texto_atual = ""
                contador_paragrafos = 0
                
                for paragrafo in doc.paragraphs:
                    texto_atual += paragrafo.text + "\n"
                    contador_paragrafos += 1
                    
                    if contador_paragrafos >= 50:
                        if texto_atual.strip():
                            textos.append(texto_atual)
                            texto_atual = ""
                            contador_paragrafos = 0
                
                if texto_atual.strip():
                    textos.append(texto_atual)
                
                if not textos:
                    textos = [""]
                
                return textos
                
            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            st.error(f"❌ Erro ao extrair texto do Word: {str(e)}")
            return []
    
    def extrair_texto_por_pagina(self, arquivo_bytes: bytes, nome_arquivo: str) -> List[str]:
        """Extrai texto baseado no tipo do arquivo"""
        tipo = self.detectar_tipo_arquivo(nome_arquivo)
        
        progress_bar = st.progress(0)
        
        try:
            if tipo == 'pdf':
                st.info("📖 Extraindo texto do PDF...")
                textos = self.extrair_texto_pdf(arquivo_bytes)
            elif tipo == 'word':
                st.info("📖 Extraindo texto do documento Word...")
                textos = self.extrair_texto_word(arquivo_bytes)
            else:
                st.error(f"❌ Tipo de arquivo não suportado: {tipo}")
                return []
            
            progress_bar.progress(1.0)
            progress_bar.empty()
            return textos
            
        except Exception as e:
            progress_bar.empty()
            st.error(f"❌ Erro ao extrair texto: {str(e)}")
            return []
    
    def normalizar_texto(self, texto: str) -> str:
        """Normaliza o texto removendo variações que não são alterações reais"""
        texto = re.sub(r'\s+', ' ', texto.strip())
        texto = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', texto)
        texto = re.sub(r'\s+([,.;:!?])', r'\1', texto)
        texto = re.sub(r'["""]', '"', texto)
        texto = re.sub(r"[''']", "'", texto)
        texto = re.sub(r'[–—]', '-', texto)
        return texto
    
    def dividir_em_paragrafos(self, texto: str) -> List[str]:
        """Divide o texto em parágrafos de forma inteligente"""
        texto = self.normalizar_texto(texto)
        paragrafos_brutos = re.split(r'\n\s*\n', texto)
        paragrafos = []
        
        for paragrafo in paragrafos_brutos:
            paragrafo = paragrafo.strip()
            if paragrafo:
                if len(paragrafo) > 500:
                    frases = re.split(r'(?<!\d)\.(?!\d)\s+', paragrafo)
                    for frase in frases:
                        frase = self.normalizar_texto(frase.strip())
                        if frase and len(frase) > 10:
                            paragrafos.append(frase)
                else:
                    paragrafo_normalizado = self.normalizar_texto(paragrafo)
                    if paragrafo_normalizado and len(paragrafo_normalizado) > 10:
                        paragrafos.append(paragrafo_normalizado)
        
        return paragrafos
    
    def calcular_similaridade(self, texto1: str, texto2: str) -> float:
        """Calcula a similaridade entre dois textos (0.0 a 1.0)"""
        if not texto1 and not texto2:
            return 1.0
        if not texto1 or not texto2:
            return 0.0
        
        texto1_norm = self.normalizar_texto(texto1)
        texto2_norm = self.normalizar_texto(texto2)
        
        matcher = difflib.SequenceMatcher(None, texto1_norm, texto2_norm)
        return matcher.ratio()
    
    def comparar_documentos(self, arquivo_ref_bytes: bytes, nome_ref: str, 
                          arquivo_novo_bytes: bytes, nome_novo: str) -> Dict:
        """Compara dois documentos e retorna as diferenças"""
        
        # Extrair textos
        textos_ref = self.extrair_texto_por_pagina(arquivo_ref_bytes, nome_ref)
        textos_novo = self.extrair_texto_por_pagina(arquivo_novo_bytes, nome_novo)
        
        if not textos_ref or not textos_novo:
            return {"erro": "Não foi possível extrair texto dos arquivos"}
        
        # Combinar textos de todas as páginas
        texto_completo_ref = "\n".join(textos_ref)
        texto_completo_novo = "\n".join(textos_novo)
        
        # Dividir em parágrafos
        paragrafos_ref = self.dividir_em_paragrafos(texto_completo_ref)
        paragrafos_novo = self.dividir_em_paragrafos(texto_completo_novo)
        
        # Encontrar diferenças
        diferencas = []
        
        # Usar SequenceMatcher para encontrar diferenças
        matcher = difflib.SequenceMatcher(None, paragrafos_ref, paragrafos_novo)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'delete':
                for i in range(i1, i2):
                    diferencas.append({
                        'tipo': 'removido',
                        'conteudo': paragrafos_ref[i],
                        'posicao_ref': i + 1,
                        'posicao_novo': None
                    })
            elif tag == 'insert':
                for j in range(j1, j2):
                    diferencas.append({
                        'tipo': 'adicionado',
                        'conteudo': paragrafos_novo[j],
                        'posicao_ref': None,
                        'posicao_novo': j + 1
                    })
            elif tag == 'replace':
                for i, j in zip(range(i1, i2), range(j1, j2)):
                    similaridade = self.calcular_similaridade(paragrafos_ref[i], paragrafos_novo[j])
                    if similaridade < 0.8:  # Considerado modificação significativa
                        diferencas.append({
                            'tipo': 'modificado',
                            'conteudo_original': paragrafos_ref[i],
                            'conteudo_novo': paragrafos_novo[j],
                            'posicao_ref': i + 1,
                            'posicao_novo': j + 1,
                            'similaridade': similaridade
                        })
        
        # Estatísticas
        total_paragrafos_ref = len(paragrafos_ref)
        total_paragrafos_novo = len(paragrafos_novo)
        total_diferencas = len(diferencas)
        
        removidos = len([d for d in diferencas if d['tipo'] == 'removido'])
        adicionados = len([d for d in diferencas if d['tipo'] == 'adicionado'])
        modificados = len([d for d in diferencas if d['tipo'] == 'modificado'])
        
        return {
            'diferencas': diferencas,
            'estatisticas': {
                'total_paragrafos_ref': total_paragrafos_ref,
                'total_paragrafos_novo': total_paragrafos_novo,
                'total_diferencas': total_diferencas,
                'removidos': removidos,
                'adicionados': adicionados,
                'modificados': modificados
            },
            'arquivos': {
                'referencia': nome_ref,
                'novo': nome_novo
            }
        }

# Classe para análise FRE
class FREAnalyzer:
    def __init__(self, api_key):
        self.client = openai.OpenAI(api_key=api_key)
        
    def extract_text_from_pdf(self, pdf_file):
        """Extrai texto de arquivo PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Erro ao extrair texto do PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extrai texto de arquivo Word"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Erro ao extrair texto do Word: {str(e)}")
            return ""
    
    def extract_text_from_file(self, uploaded_file):
        """Extrai texto baseado no tipo de arquivo"""
        if uploaded_file.type == "application/pdf":
            return self.extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                   "application/msword"]:
            return self.extract_text_from_docx(uploaded_file)
        else:
            st.error("Formato de arquivo não suportado. Use PDF ou Word.")
            return ""
    
    def analyze_fre_section(self, fre_text, cvm_references, section_name, section_content):
        """Analisa uma seção específica do FRE contra as normas CVM"""
        
        prompt = f"""
        Você é um especialista em regulamentação CVM e análise de Formulários de Referência (FRE).
        
        Analise a seção "{section_name}" do FRE fornecido contra as normas e orientações CVM.
        
        SEÇÃO ANALISADA:
        {section_content[:3000]}...
        
        NORMAS CVM DE REFERÊNCIA:
        {cvm_references[:5000]}...
        
        Para esta seção, identifique:
        
        1. CONFORMIDADE: Está em conformidade com as normas CVM?
        2. COMPLETUDE: Todas as informações obrigatórias estão presentes?
        3. QUALIDADE: A informação está clara, objetiva e completa?
        4. PONTOS DE ATENÇÃO: Identifique problemas específicos
        5. SUGESTÕES: Recomendações de melhoria com citação obrigatória dos artigos CVM
        
        RESPONDA EM JSON com esta estrutura:
        {{
            "secao": "{section_name}",
            "conformidade": "CONFORME/NAO_CONFORME/PARCIALMENTE_CONFORME",
            "criticidade": "CRITICO/ATENCAO/SUGESTAO",
            "pontos_atencao": [
                {{
                    "problema": "descrição do problema",
                    "criticidade": "CRITICO/ATENCAO/SUGESTAO",
                    "artigo_cvm": "artigo específico da norma CVM",
                    "sugestao": "recomendação específica de melhoria"
                }}
            ],
            "resumo": "resumo geral da análise desta seção"
        }}
        
        IMPORTANTE: 
        - Cite OBRIGATORIAMENTE os artigos específicos das normas CVM
        - Use criticidade CRITICO para não conformidades graves
        - Use ATENCAO para informações incompletas
        - Use SUGESTAO para melhorias recomendadas
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=2000
            )
            
            result = response.choices[0].message.content
            
            try:
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                json_str = result[json_start:json_end]
                return json.loads(json_str)
            except:
                return {
                    "secao": section_name,
                    "conformidade": "ERRO_ANALISE",
                    "criticidade": "ATENCAO",
                    "pontos_atencao": [{
                        "problema": "Erro na análise automática",
                        "criticidade": "ATENCAO",
                        "artigo_cvm": "Resolução CVM nº 80/22",
                        "sugestao": "Revisar manualmente esta seção"
                    }],
                    "resumo": "Erro na análise automática desta seção"
                }
                
        except Exception as e:
            st.error(f"Erro na análise da seção {section_name}: {str(e)}")
            return None
    
    def extract_fre_sections(self, fre_text):
        """Extrai as seções principais do FRE"""
        sections = {}
        
        section_patterns = [
            r"1\.1\s+Histórico do emissor",
            r"1\.2\s+Descrição das principais atividades",
            r"1\.3\s+Informações relacionadas aos segmentos operacionais",
            r"1\.4\s+Produção/Comercialização/Mercados",
            r"1\.5\s+Principais clientes",
            r"1\.6\s+Efeitos relevantes da regulação estatal",
            r"1\.9\s+Informações ambientais sociais e de governança",
            r"2\.1\s+Condições financeiras e patrimoniais",
            r"2\.2\s+Resultados operacional e financeiro",
            r"4\.1\s+Descrição dos fatores de risco",
            r"7\.1\s+Principais características dos órgãos de administração",
            r"8\.1\s+Política ou prática de remuneração",
            r"11\.1\s+Regras, políticas e práticas",
            r"12\.1\s+Informações sobre o capital social"
        ]
        
        lines = fre_text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            section_found = False
            for pattern in section_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    current_section = line.strip()
                    current_content = [line]
                    section_found = True
                    break
            
            if not section_found and current_section:
                current_content.append(line)
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections

# Função principal
def main():
    render_header()
    
    # Configuração da API Key na sidebar
    with st.sidebar:
        st.markdown("### ⚙️ Configurações")
        api_key = st.text_input(
            "Chave API OpenAI",
            type="password",
            value=st.session_state.api_key,
            help="Necessária apenas para a análise FRE vs CVM"
        )
        
        if api_key != st.session_state.api_key:
            st.session_state.api_key = api_key
        
        st.markdown("---")
        st.markdown("### 📋 Sobre a Plataforma")
        st.markdown("""
        **Comparador de Documentos:**
        - Compara PDFs e documentos Word
        - Identifica diferenças reais de conteúdo
        - Gera relatórios detalhados
        
        **Analisador FRE vs CVM:**
        - Analisa Formulários de Referência
        - Verifica conformidade com normas CVM
        - Usa inteligência artificial (GPT-4)
        """)
    
    # Abas principais
    tab1, tab2 = st.tabs(["📊 Comparador de Documentos", "📋 Analisador FRE vs CVM"])
    
    with tab1:
        render_document_comparator()
    
    with tab2:
        render_fre_analyzer()

def render_document_comparator():
    """Renderiza a interface do comparador de documentos"""
    st.markdown("### 📊 Comparação de Documentos")
    st.markdown("Compare dois arquivos (PDF ou Word) e identifique as diferenças de conteúdo.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 📄 Arquivo de Referência")
        arquivo_ref = st.file_uploader(
            "Selecione o arquivo de referência",
            type=['pdf', 'docx', 'doc'],
            key="ref_file"
        )
    
    with col2:
        st.markdown("#### 📄 Arquivo Novo")
        arquivo_novo = st.file_uploader(
            "Selecione o arquivo novo",
            type=['pdf', 'docx', 'doc'],
            key="new_file"
        )
    
    if arquivo_ref and arquivo_novo:
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col2:
            if st.button("🔍 Comparar Documentos", use_container_width=True):
                with st.spinner("Comparando documentos..."):
                    comparator = DocumentComparator()
                    
                    # Validar arquivos
                    if not comparator.validar_arquivo(arquivo_ref.read(), arquivo_ref.name):
                        return
                    arquivo_ref.seek(0)
                    
                    if not comparator.validar_arquivo(arquivo_novo.read(), arquivo_novo.name):
                        return
                    arquivo_novo.seek(0)
                    
                    # Comparar
                    resultado = comparator.comparar_documentos(
                        arquivo_ref.read(), arquivo_ref.name,
                        arquivo_novo.read(), arquivo_novo.name
                    )
                    
                    if "erro" in resultado:
                        st.error(f"❌ {resultado['erro']}")
                        return
                    
                    st.session_state.comparison_results = resultado
                    st.success("✅ Comparação concluída!")
    
    # Exibir resultados
    if st.session_state.comparison_results:
        resultado = st.session_state.comparison_results
        
        st.markdown("---")
        st.markdown("### 📈 Resultados da Comparação")
        
        # Estatísticas
        stats = resultado['estatisticas']
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{stats['total_diferencas']}</div>
                <div class="metric-label">Total de Diferenças</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{stats['adicionados']}</div>
                <div class="metric-label">Parágrafos Adicionados</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{stats['removidos']}</div>
                <div class="metric-label">Parágrafos Removidos</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{stats['modificados']}</div>
                <div class="metric-label">Parágrafos Modificados</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Filtros
        st.markdown("### 🔍 Filtros")
        col1, col2 = st.columns(2)
        
        with col1:
            tipos_filtro = st.multiselect(
                "Tipos de diferença",
                ["adicionado", "removido", "modificado"],
                default=["adicionado", "removido", "modificado"]
            )
        
        with col2:
            mostrar_detalhes = st.checkbox("Mostrar detalhes completos", value=True)
        
        # Exibir diferenças
        if tipos_filtro:
            diferencas_filtradas = [d for d in resultado['diferencas'] if d['tipo'] in tipos_filtro]
            
            st.markdown(f"### 📋 Diferenças Encontradas ({len(diferencas_filtradas)})")
            
            for i, diff in enumerate(diferencas_filtradas):
                if diff['tipo'] == 'adicionado':
                    st.markdown(f"""
                    <div class="paragrafo-adicionado">
                        <strong>➕ Parágrafo Adicionado (Posição: {diff['posicao_novo']})</strong><br>
                        {diff['conteudo'][:500]}{'...' if len(diff['conteudo']) > 500 and not mostrar_detalhes else diff['conteudo']}
                    </div>
                    """, unsafe_allow_html=True)
                
                elif diff['tipo'] == 'removido':
                    st.markdown(f"""
                    <div class="paragrafo-removido">
                        <strong>➖ Parágrafo Removido (Posição: {diff['posicao_ref']})</strong><br>
                        {diff['conteudo'][:500]}{'...' if len(diff['conteudo']) > 500 and not mostrar_detalhes else diff['conteudo']}
                    </div>
                    """, unsafe_allow_html=True)
                
                elif diff['tipo'] == 'modificado':
                    st.markdown(f"""
                    <div class="paragrafo-modificado">
                        <strong>🔄 Parágrafo Modificado (Ref: {diff['posicao_ref']}, Novo: {diff['posicao_novo']})</strong><br>
                        <strong>Original:</strong> {diff['conteudo_original'][:300]}{'...' if len(diff['conteudo_original']) > 300 and not mostrar_detalhes else diff['conteudo_original']}<br>
                        <strong>Novo:</strong> {diff['conteudo_novo'][:300]}{'...' if len(diff['conteudo_novo']) > 300 and not mostrar_detalhes else diff['conteudo_novo']}
                    </div>
                    """, unsafe_allow_html=True)

def render_fre_analyzer():
    """Renderiza a interface do analisador FRE"""
    st.markdown("### 📋 Análise FRE vs Normas CVM")
    st.markdown("Analise Formulários de Referência contra as normas da CVM usando inteligência artificial.")
    
    if not st.session_state.api_key:
        st.markdown("""
        <div class="warning-box">
            ⚠️ <strong>Chave API OpenAI necessária</strong><br>
            Configure sua chave API OpenAI na barra lateral para usar esta funcionalidade.
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Upload do arquivo FRE
    st.markdown("#### 📄 Upload do Formulário de Referência")
    fre_file = st.file_uploader(
        "Selecione o arquivo FRE (PDF ou Word)",
        type=['pdf', 'docx', 'doc'],
        key="fre_file"
    )
    
    if fre_file:
        st.session_state.fre_filename = fre_file.name
        
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col2:
            if st.button("🔍 Analisar FRE", use_container_width=True):
                with st.spinner("Analisando FRE contra normas CVM..."):
                    try:
                        analyzer = FREAnalyzer(st.session_state.api_key)
                        
                        # Extrair texto do FRE
                        fre_text = analyzer.extract_text_from_file(fre_file)
                        
                        if not fre_text:
                            st.error("❌ Não foi possível extrair texto do arquivo FRE")
                            return
                        
                        # Extrair seções do FRE
                        sections = analyzer.extract_fre_sections(fre_text)
                        
                        if not sections:
                            st.error("❌ Não foi possível identificar seções no FRE")
                            return
                        
                        # Normas CVM de referência (simplificado)
                        cvm_references = """
                        Resolução CVM nº 80/22 - Formulário de Referência
                        
                        Art. 1º - O formulário de referência deve conter informações sobre:
                        I - histórico do emissor;
                        II - atividades do emissor;
                        III - fatores de risco;
                        IV - dados econômico-financeiros;
                        V - comentários dos diretores;
                        VI - recursos humanos;
                        VII - controladores e administradores;
                        VIII - remuneração dos administradores;
                        IX - partes relacionadas;
                        X - contratos relevantes;
                        XI - governança corporativa;
                        XII - capital social;
                        XIII - valores mobiliários;
                        XIV - distribuição de dividendos;
                        XV - política de negociação;
                        XVI - informações sobre aquisições.
                        
                        As informações devem ser claras, precisas, verdadeiras, atuais, completas e, quando necessário, quantificadas.
                        """
                        
                        # Analisar cada seção
                        results = []
                        progress_bar = st.progress(0)
                        
                        for i, (section_name, section_content) in enumerate(sections.items()):
                            progress_bar.progress((i + 1) / len(sections))
                            
                            analysis = analyzer.analyze_fre_section(
                                fre_text, cvm_references, section_name, section_content
                            )
                            
                            if analysis:
                                results.append(analysis)
                            
                            time.sleep(1)  # Evitar rate limiting
                        
                        progress_bar.empty()
                        
                        st.session_state.analysis_results = results
                        st.session_state.analysis_completed = True
                        st.success("✅ Análise concluída!")
                        
                    except Exception as e:
                        st.error(f"❌ Erro na análise: {str(e)}")
    
    # Exibir resultados da análise
    if st.session_state.analysis_results and st.session_state.analysis_completed:
        results = st.session_state.analysis_results
        
        st.markdown("---")
        st.markdown("### 📊 Resultados da Análise")
        
        # Estatísticas gerais
        total_secoes = len(results)
        conformes = len([r for r in results if r.get('conformidade') == 'CONFORME'])
        nao_conformes = len([r for r in results if r.get('conformidade') == 'NAO_CONFORME'])
        parcialmente_conformes = len([r for r in results if r.get('conformidade') == 'PARCIALMENTE_CONFORME'])
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value">{total_secoes}</div>
                <div class="metric-label">Seções Analisadas</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value" style="color: #22c55e;">{conformes}</div>
                <div class="metric-label">Conformes</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value" style="color: #f59e0b;">{parcialmente_conformes}</div>
                <div class="metric-label">Parcialmente Conformes</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown(f"""
            <div class="metric-container">
                <div class="metric-value" style="color: #ef4444;">{nao_conformes}</div>
                <div class="metric-label">Não Conformes</div>
            </div>
            """, unsafe_allow_html=True)
        
        # Filtros
        st.markdown("### 🔍 Filtros")
        col1, col2 = st.columns(2)
        
        with col1:
            conformidade_filtro = st.multiselect(
                "Status de Conformidade",
                ["CONFORME", "PARCIALMENTE_CONFORME", "NAO_CONFORME"],
                default=["CONFORME", "PARCIALMENTE_CONFORME", "NAO_CONFORME"]
            )
        
        with col2:
            criticidade_filtro = st.multiselect(
                "Nível de Criticidade",
                ["CRITICO", "ATENCAO", "SUGESTAO"],
                default=["CRITICO", "ATENCAO", "SUGESTAO"]
            )
        
        # Exibir análises detalhadas
        results_filtrados = [
            r for r in results 
            if r.get('conformidade') in conformidade_filtro 
            and r.get('criticidade') in criticidade_filtro
        ]
        
        st.markdown(f"### 📋 Análise Detalhada ({len(results_filtrados)} seções)")
        
        for result in results_filtrados:
            # Cor baseada na conformidade
            if result.get('conformidade') == 'CONFORME':
                cor_header = "#22c55e"
            elif result.get('conformidade') == 'PARCIALMENTE_CONFORME':
                cor_header = "#f59e0b"
            else:
                cor_header = "#ef4444"
            
            with st.expander(f"📄 {result.get('secao', 'Seção desconhecida')}", expanded=False):
                st.markdown(f"""
                <div style="background-color: {cor_header}20; border-left: 4px solid {cor_header}; padding: 15px; border-radius: 5px; margin: 10px 0;">
                    <strong>Status:</strong> {result.get('conformidade', 'N/A')}<br>
                    <strong>Criticidade:</strong> {result.get('criticidade', 'N/A')}<br>
                    <strong>Resumo:</strong> {result.get('resumo', 'N/A')}
                </div>
                """, unsafe_allow_html=True)
                
                if result.get('pontos_atencao'):
                    st.markdown("**Pontos de Atenção:**")
                    for ponto in result['pontos_atencao']:
                        criticidade_cor = {
                            'CRITICO': '#ef4444',
                            'ATENCAO': '#f59e0b',
                            'SUGESTAO': '#3b82f6'
                        }.get(ponto.get('criticidade', 'SUGESTAO'), '#3b82f6')
                        
                        st.markdown(f"""
                        <div style="background-color: {criticidade_cor}15; border-left: 3px solid {criticidade_cor}; padding: 10px; margin: 5px 0; border-radius: 3px;">
                            <strong>Problema:</strong> {ponto.get('problema', 'N/A')}<br>
                            <strong>Artigo CVM:</strong> {ponto.get('artigo_cvm', 'N/A')}<br>
                            <strong>Sugestão:</strong> {ponto.get('sugestao', 'N/A')}
                        </div>
                        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()

